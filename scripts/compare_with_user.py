"""Compare our updated file with user's reference file on all 7 changes."""
from docx import Document

our = Document('/home/z/my-project/YMK/KTP/МДК_05_02_КТП.docx')
usr = Document('/home/z/my-project/YMK/KTP/КТП МДК 05.02 М2 курс.docx')

print('=== 1. Ministry ===')
print(f'  Our:  {our.paragraphs[0].text}')
print(f'  User: {usr.paragraphs[0].text}')
print(f'  MATCH: {our.paragraphs[0].text == usr.paragraphs[0].text}')

print('\n=== 2. KTP font size ===')
for d, label in [(our, 'Our'), (usr, 'User')]:
    for p in d.paragraphs:
        if 'КАЛЕНДАРНО' in p.text:
            print(f'  {label}: {p.runs[0].font.size} style={p.style.name}')
            break

print('\n=== 3. Table 1 structure ===')
for d, label in [(our, 'Our'), (usr, 'User')]:
    t = d.tables[1]
    print(f'  {label}: {len(t.rows)} rows x {len(t.columns)} cols')
    # Check row 3 content
    texts = [c.text[:25] for c in t.rows[3].cells]
    print(f'    R[3]: {texts[:6]}...')

print('\n=== 4. Bardakov ===')
for d, label in [(our, 'Our'), (usr, 'User')]:
    t = d.tables[2]
    print(f'  {label}: "{t.rows[4].cells[10].text}"')

print('\n=== 5. Control column sample ===')
for d, label in [(our, 'Our'), (usr, 'User')]:
    t = d.tables[2]
    print(f'  {label} R[7]: "{t.rows[7].cells[9].text}"')

print('\n=== 6. Diff rows ===')
for d, label in [(our, 'Our'), (usr, 'User')]:
    t = d.tables[2]
    diff_rows = [(ri, r.cells[0].text, r.cells[1].text[:30], r.cells[4].text) 
                 for ri, r in enumerate(t.rows) if chr(1044)+chr(1080)+chr(1092) in r.cells[1].text]
    print(f'  {label}: {diff_rows}')

print('\n=== 7. DI table ===')
for d, label in [(our, 'Our'), (usr, 'User')]:
    t = d.tables[5]
    print(f'  {label}: {len(t.rows)} rows')
    for ri, row in enumerate(t.rows):
        print(f'    R[{ri}]: {[c.text[:35] for c in row.cells]}')

print('\n=== 7b. EI table ===')
for d, label in [(our, 'Our'), (usr, 'User')]:
    t = d.tables[6]
    print(f'  {label}: {len(t.rows)} rows x {len(t.columns)} cols')
    for ri, row in enumerate(t.rows):
        print(f'    R[{ri}]: {[c.text[:35] for c in row.cells]}')
