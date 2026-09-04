#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Валидатор сгенерированного КТП по разделу 12 KTP_INSTRUCTIONS.md.

Использование:
    python3 scripts/validate_ktp.py --doc "KTP/КТП МДК 05.02 М2 курс.docx" --data data.json

Проверки:
  1.  Валидный ZIP (OOXML)
  2.  7 таблиц: УТВЕРЖДАЮ, Т1, Т2, 2а, 2б, 2в, 2г
  3.  Т1: 12 колонок, сумма gridSpan каждой строки = 12
  4.  Т2: 11 колонок, сумма gridSpan каждой строки = 11
  5.  ЭИ (2г): 5 колонок (4 + URL)
  6.  Шапка Т2 и название министерства — строго как в эталоне
  7.  Сумма часов занятий = общий объём; сумма практики = плановой;
      титул: объём и «в т.ч. практическая подготовка» = РП (табл. 3.1)
  8.  Сквозная нумерация занятий без пропусков
  9.  Строка диф.зачёта/аттестации: ОК/ПК/вид/обеспечение/задания/контроль пустые
  10. Литература: в ОИ нет книг старше 5 лет; ЭИ имеют URL
  11. КТИЧЕСКИЕ тексты: без заливки ячеек (кроме белой), шрифт Times New Roman
"""
import argparse
import json
import re
import sys
import zipfile
from collections import Counter

from docx import Document

W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

MINISTRY = 'МИНИСТЕРСТВО ОБРАЗОВАНИЯ РОСТОВСКОЙ ОБЛАСТИ'
T2_HEADERS = {
    '№ занятия', 'Наименование разделов\nпрофессионального модуля,\nтем и занятий по МДК',
    'Обязательная учебная нагрузка', 'Коды формируемых компетенции',
    'Кол-во\nчасов', 'В форме практической подготовки', 'Вид занятия',
    'Материальное и информационное\nобеспечение занятий\n', 'Задания для студентов\n',
    'Формы и методы контроля', 'ФИО преподавателя \n', 'ОК', 'ПК',
}

results = []


def check(name, ok, detail=''):
    results.append((name, ok, detail))


def grid_cols(row):
    total = 0
    for tc in row.findall(W + 'tc'):
        tcPr = tc.find(W + 'tcPr')
        gs = 1
        if tcPr is not None:
            gs_el = tcPr.find(W + 'gridSpan')
            if gs_el is not None:
                gs = int(gs_el.get(W + 'val'))
        total += gs
    return total


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--doc', required=True)
    ap.add_argument('--data', required=True)
    ap.add_argument('--year', type=int, default=2026, help='текущий год для правила 5 лет')
    args = ap.parse_args()

    with open(args.data, encoding='utf-8') as f:
        data = json.load(f)

    # 1. ZIP
    try:
        bad = zipfile.ZipFile(args.doc).testzip()
        check('1. Валидный ZIP-архив', bad is None, str(bad) if bad else '')
    except Exception as e:
        check('1. Валидный ZIP-архив', False, str(e))

    doc = Document(args.doc)
    tables = doc.tables

    # 2. Количество таблиц
    check('2. Ровно 7 таблиц (УТВЕРЖДАЮ, Т1, Т2, 2а, 2б, 2в, 2г)',
          len(tables) == 7, f'найдено {len(tables)}')

    if len(tables) >= 7:
        t1, t2 = tables[1], tables[2]
        t_ei = tables[6]

        # 3. Т1 — 12 колонок
        ok = len(t1.columns) == 12
        bad_rows = [i for i, r in enumerate(t1._tbl.findall(W + 'tr')) if grid_cols(r) != 12]
        check('3. Таблица 1: 12 колонок, все строки = 12', ok and not bad_rows,
              f'плохие строки: {bad_rows}' if bad_rows else f'{len(t1.rows)} строк')

        # 4. Т2 — 11 колонок
        ok = len(t2.columns) == 11
        bad_rows = [i for i, r in enumerate(t2._tbl.findall(W + 'tr')) if grid_cols(r) != 11]
        check('4. Таблица 2: 11 колонок, все строки = 11', ok and not bad_rows,
              f'плохие строки: {bad_rows}' if bad_rows else f'{len(t2.rows)} строк')

        # 5. ЭИ — 5 колонок
        check('5. Таблица 2г (ЭИ): 5 колонок (вкл. URL)', len(t_ei.columns) == 5,
              f'{len(t_ei.columns)} колонок')
        has_url_hdr = 'URL' in t_ei.cell(0, 4).text
        check('   и заголовок 5-й колонки — URL (ссылка)', has_url_hdr,
              t_ei.cell(0, 4).text.strip())

        # 6. Шапка Т2 и министерство
        hdr_cells = set()
        for ri in range(3):
            for ci in range(11):
                hdr_cells.add(t2.cell(ri, ci).text)
        missing = {h for h in T2_HEADERS if h not in hdr_cells}
        check('6. Тексты шапки Т2 соответствуют эталону', not missing,
              f'нет: {missing}' if missing else '')
        ministry_found = any(p.text.strip() == MINISTRY for p in doc.paragraphs)
        check('   Название министерства строго по эталону', ministry_found)

        # 7. Суммы часов
        nums, sum_h, sum_p = [], 0, 0
        att_rows = []
        theme_hdr, theme_sum, cur_theme = {}, {}, None  # построчная сверка с РП 3.2 (контрольные суммы)
        for ri in range(4, len(t2.rows)):
            c0 = t2.cell(ri, 0).text.strip()
            c1 = t2.cell(ri, 1).text.strip()
            if c1.startswith('Итого'):
                break
            m_theme = re.match(r'Тема\s*(\d+\.\d)', c1)
            if m_theme:
                cur_theme = m_theme.group(1)
            if not c0:
                # строка-заголовок темы: кол.4 = плановая практика темы (по РП 3.2)
                if m_theme:
                    p_hdr = t2.cell(ri, 3).text.strip()
                    theme_hdr[cur_theme] = int(p_hdr) if p_hdr else 0
                continue
            nums.append(int(c0))
            h = t2.cell(ri, 2).text.strip()
            p = t2.cell(ri, 3).text.strip()
            sum_h += int(h) if h else 0
            sum_p += int(p) if p else 0
            if cur_theme:
                theme_sum[cur_theme] = theme_sum.get(cur_theme, 0) + (int(p) if p else 0)
            if 'зачет' in c1.lower() or 'зачёт' in c1.lower() or 'экзамен' in c1.lower():
                att_rows.append((ri, c1, h, p,
                                 t2.cell(ri, 4).text.strip(), t2.cell(ri, 5).text.strip(),
                                 t2.cell(ri, 6).text.strip(), t2.cell(ri, 7).text.strip(),
                                 t2.cell(ri, 8).text.strip(), t2.cell(ri, 9).text.strip()))
        check('7а. Σ часов занятий = общему объёму МДК',
              sum_h == data['total_hours'], f'{sum_h} != {data["total_hours"]}' if sum_h != data['total_hours'] else f'{sum_h} ч')
        check('7б. Σ практики (строки занятий) = практике МДК по РП, табл. 3.2',
              sum_p == data['practice_hours'], f'{sum_p} != {data["practice_hours"]}' if sum_p != data['practice_hours'] else f'{sum_p} ч')
        # 7ж/7з — построчные контрольные суммы (перенос из РП 3.2):
        bad_themes = {k: (theme_hdr.get(k, 0), theme_sum.get(k, 0))
                      for k in set(theme_hdr) | set(theme_sum)
                      if theme_hdr.get(k, 0) != theme_sum.get(k, 0)}
        check('7ж. По КАЖДОЙ теме: Σ практики занятий = значению заголовка темы (построчный перенос из РП 3.2)',
              bool(theme_hdr) and not bad_themes,
              f'несходятся: {bad_themes}' if bad_themes else f'{len(theme_hdr)} тем ✓')
        check('7з. Σ кол.4 заголовков тем = практике МДК (дублирует Σ по занятиям)',
              sum(theme_hdr.values()) == data['practice_hours'],
              f'{sum(theme_hdr.values())} != {data["practice_hours"]}'
              if sum(theme_hdr.values()) != data['practice_hours'] else f'{sum(theme_hdr.values())} ч')
        th = data.get('theory_hours', 0)
        att_h = data.get('attestation', {}).get('hours', 0)
        check('7в. Теория + практика + аттестация = общий объём',
              th + data['practice_hours'] + att_h == data['total_hours'],
              f'{th}+{data["practice_hours"]}+{att_h} != {data["total_hours"]}'
              if th + data['practice_hours'] + att_h != data['total_hours'] else f'{th}+{data["practice_hours"]}+{att_h}')

        # 7г-7е. Титул: часы (источник — кол. 3/4 строки МДК тематического плана РП)
        title = '\n'.join(p.text for p in doc.paragraphs)
        m = re.search(r'Объем образовательной программы:\s*_{0,20}(\d+)_{0,20}\s*\((час\w*)\)', title)
        check('7г. Титул: объём образовательной программы = total_hours (кол. 3 строки МДК РП)',
              bool(m) and int(m.group(1)) == data['total_hours'],
              (m.group(1) if m else 'строка не найдена') + f' != {data["total_hours"]}'
              if not m or int(m.group(1)) != data['total_hours'] else f'{data["total_hours"]} ч')
        m2 = re.search(r'в том числе в форме практической подготовки\s*_{0,20}(\d+)_{0,20}\s*\((час\w*)\)', title)
        check('7д. Титул: «в т.ч. в форме практической подготовки» = practice_hours (кол. 4 строки МДК РП)',
              bool(m2) and int(m2.group(1)) == data['practice_hours'],
              (m2.group(1) if m2 else 'строка не найдена') + f' != {data["practice_hours"]}'
              if not m2 or int(m2.group(1)) != data['practice_hours'] else f'{data["practice_hours"]} ч')
        # 7е. Склонение «час» в титуле: 1→час, 2-4→часа, 5-20→часов (8→часов, 2→часа)
        def hour_word(n):
            if n in (None, '', 0):
                return 'часов'
            n = int(n)
            if n % 100 in (11, 12, 13, 14):
                return 'часов'
            if n % 10 == 1:
                return 'час'
            if n % 10 in (2, 3, 4):
                return 'часа'
            return 'часов'
        pairs = re.findall(r'(\d+)\s*_*\s*\((часа|часов|час)\)', title)
        bad_decl = [(n, w) for n, w in pairs if hour_word(int(n)) != w]
        check('7е. Склонение слова «час» в титуле соответствует числам (8→часов, 2→часа)',
              len(pairs) >= 3 and not bad_decl,
              f'несогласованы: {bad_decl}' if bad_decl else f'{len(pairs)} строк ✓')

        # 8. Сквозная нумерация
        ok = nums == list(range(1, len(nums) + 1))
        check('8. Нумерация занятий сквозная без пропусков', ok,
              f'{len(nums)} занятий, последним № {nums[-1] if nums else "-"}')

        # 9. Строка аттестации пустая в графах 4-10
        if att_rows:
            ri, name, h, p, vid, okc, pk, mat, zad, kontr = att_rows[-1]
            empty_ok = all(x == '' for x in (vid, okc, pk, mat, zad, kontr))
            check('9. Аттестация: номер/название/часы есть, графы 5-10 пусты (эталон)',
                  empty_ok, f'строка «{name}»: вид={vid!r}, ок={okc!r}, контроль={kontr!r}')

        # 10. Литература
        def year_of(s):
            m = re.search(r'(20\d{2}|19\d{2})', s or '')
            return int(m.group(1)) if m else None

        oi_bad = []
        for ri in range(1, len(tables[4].rows)):
            pub = tables[4].cell(ri, 3).text
            y = year_of(pub)
            if y and y < args.year - 5:  # старше 5 лет = издано ДО 2021 (для 2026);
                # 2021+ — свежие (KTP_INSTRUCTIONS.md, раздел 3)
                oi_bad.append((tables[4].cell(ri, 0).text.strip(), y))
        check('10а. В ОИ нет книг старше 5 лет', not oi_bad, str(oi_bad) if oi_bad else '')

        ei_no_url = [tables[6].cell(ri, 0).text.strip()
                     for ri in range(1, len(tables[6].rows))
                     if not tables[6].cell(ri, 4).text.strip().startswith(('http', 'www'))]
        check('10б. Все ЭИ имеют URL', not ei_no_url, str(ei_no_url) if ei_no_url else '')

        # 11. Заливка и шрифт
        xml = doc.element.xml
        fills = Counter(re.findall(r'w:fill="([0-9A-Fa-f]{6})"', xml))
        non_white = {k: v for k, v in fills.items() if k.upper() not in ('FFFFFF', 'AUTO')}
        check('11а. Заливки ячеек отсутствуют (строго по эталону)', not non_white, str(non_white))
        # 11б. Шрифты реальных run'ов (pPr/rPr — свойства маркера конца абзаца —
        # не отображаются и в эталоне исторически содержат Arial: не считаем ошибкой)
        run_fonts = set(re.findall(
            r'<w:r(?:\s[^>]*)?>(?:(?!</w:r>).)*?<w:rFonts w:ascii="([^"]+)"',
            xml, re.S))
        bad_fonts = {f for f in run_fonts if 'Times' not in f}
        check('11б. Весь видимый текст Times New Roman', not bad_fonts, str(bad_fonts))

        # 12. Таблица 1 и высоты строк (по исправленному образцу эталона 05.02)
        # 12а. Строка «Компл. дифф. зачет»: часы в графе 6, остальные графы пустые
        att_t1 = None
        for ri in range(4, len(t1.rows)):
            c0 = t1.cell(ri, 0).text.strip()
            if c0.lower().startswith('компл'):
                att_t1 = [t1.cell(ri, c).text.strip() for c in range(12)]
                break
        ok = (att_t1 is not None
              and att_t1[0].lower().startswith('компл')
              and att_t1[5] == str(data.get('attestation', {}).get('hours', 2))
              and all(att_t1[c] == '' for c in (1, 2, 3, 4, 6, 7, 8, 9, 10, 11)))
        check('12а. Т1: строка «Компл. дифф. зачет» — часы в графе 6, остальные графы пусты (эталон)',
              ok, str(att_t1) if not ok else '')
        # 12б. Суммы Т1: Σ объёмов семестров = Всего (КДЗ внутри последнего семестра);
        #      графа «Теоретические занятия» Всего = Σ теорий семестров + КДЗ; практика = РП
        if att_t1 is not None:
            sems_t1 = []
            for ri in range(5, len(t1.rows)):
                c0 = t1.cell(ri, 0).text.strip()
                if c0 == 'МДК ' + data['mdk_code']:
                    sems_t1.append([t1.cell(ri, c).text.strip() for c in range(12)])
            tot_row = None
            for ri in range(len(t1.rows) - 1, 4, -1):
                if t1.cell(ri, 0).text.strip() == 'Всего':
                    tot_row = [t1.cell(ri, c).text.strip() for c in range(12)]
                    break
            att_h = data.get('attestation', {}).get('hours', 0)
            ok_tot = (sum(int(r[3]) for r in sems_t1) == data['total_hours']
                      and tot_row is not None and int(tot_row[3]) == data['total_hours'])
            check('12б-1. Т1: Σ объёмов семестров = Всего = total_hours (КДЗ внутри последнего семестра)',
                  ok_tot, f'семестры {[r[3] for r in sems_t1]} vs {tot_row[3] if tot_row else "-"}')
            ok_th = tot_row is not None and int(tot_row[5]) == (
                sum(int(r[5]) for r in sems_t1) + att_h)
            check('12б-2. Т1: теория Всего = Σ теорий семестров + КДЗ',
                  ok_th, f'{tot_row[5] if tot_row else "-"} != {sum(int(r[5]) for r in sems_t1)}+{att_h}'
                  if not ok_th else '')
            ok_pr = tot_row is not None and int(tot_row[7]) == data['practice_hours']
            check('12б-3. Т1: практика Всего = практике МДК по РП', ok_pr,
                  f'{tot_row[7] if tot_row else "-"} != {data["practice_hours"]}' if not ok_pr else '')
        # 12в. Высоты строк: Т2 — шапка 20/230/1695, все строки с trHeight;
        #      Т1 — шапка 219/234/1758, номерная 237, данные 215
        def tr_heights(tbl):
            out = []
            for tr in tbl._tbl.findall(W + 'tr'):
                th = tr.find(W + 'trPr/' + W + 'trHeight')
                out.append(th.get(W + 'val') if th is not None else None)
            return out
        h2 = tr_heights(t2)
        ok = (len(h2) >= 4 and h2[0] == '20' and h2[1] == '230' and h2[2] == '1695'
              and all(v is not None for v in h2))
        check('12в-1. Т2: высоты строк по эталону (20/230/1695, все строки заданы)',
              ok, f'первые 4: {h2[:4]}, без высоты: {[i for i, v in enumerate(h2) if v is None][:5]}')
        h1 = tr_heights(t1)
        ok = (len(h1) >= 12 and h1[1] == '219' and h1[2] == '234' and h1[3] == '1758'
              and h1[4] == '237' and all(v == '215' for v in h1[5:] if v is not None))
        check('12в-2. Т1: высоты строк по эталону (219/234/1758/237, данные 215)',
              ok, f'первые 6: {h1[:6]}')

    # Итог
    print('=' * 72)
    failed = 0
    for name, ok, detail in results:
        mark = '✅' if ok else '❌'
        if not ok:
            failed += 1
        print(f'{mark} {name}' + (f'  — {detail}' if detail and not ok else
                                   (f'  ({detail})' if detail else '')))
    print('=' * 72)
    print('ИТОГО: %d проверок, провалено: %d' % (len(results), failed))
    sys.exit(1 if failed else 0)


if __name__ == '__main__':
    main()
