"""Apply 7 manual changes from user's template to our generated KTP."""
from docx import Document
from docx.shared import Pt, Emu, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.oxml.ns import qn
from copy import deepcopy
import lxml.etree as ET
import re, sys

DOC_PATH = '/home/z/my-project/YMK/KTP/МДК_05_02_КТП.docx'
USER_PATH = '/home/z/my-project/YMK/KTP/КТП МДК 05.02 М2 курс.docx'

doc = Document(DOC_PATH)
user_doc = Document(USER_PATH)

# ---- Copy Заголовок1 style from user file ----
def ensure_style_zagolovok1(target_doc, source_doc):
    try:
        target_doc.styles['Заголовок1']
        return
    except KeyError:
        pass
    for s in source_doc.styles.element.findall(qn('w:style')):
        name_el = s.find(qn('w:name'))
        if name_el is not None and name_el.get(qn('w:val')) == 'Заголовок1':
            target_doc.styles.element.append(deepcopy(s))
            print('Copied Заголовок1 style')
            return

ensure_style_zagolovok1(doc, user_doc)

# ---- Helpers ----
def srf(run, size_pt=12, bold=None, underline=None, name='Times New Roman'):
    """Set run font properties."""
    run.font.size = Pt(size_pt)
    run.font.name = name
    if bold is not None:
        run.bold = bold
    if underline is not None:
        run.font.underline = underline
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = ET.SubElement(run._element, qn('w:rPr'))
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = ET.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:eastAsia'), name)

def sc(row, ci, text, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER):
    """Set cell text with formatting."""
    cell = row.cells[ci]
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    run.bold = bold

def find_para(doc, substr):
    for i, p in enumerate(doc.paragraphs):
        if substr in p.text:
            return i, p
    return None, None

# ============================================================
# ИЗМЕНЕНИЯ 1-2: Титул + Министерство
# ============================================================
# P[0]: Ministry name (shortened)
p = doc.paragraphs[0]
p.style = doc.styles['Заголовок1']
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.clear()
srf(p.add_run('МИНИСТЕРСТВО ОБРАЗОВАНИЯ РОСТОВСКОЙ ОБЛАСТИ'), 12, bold=True)

# P[2-3]: Institution lines
for pi in [2, 3]:
    p = doc.paragraphs[pi]
    p.style = doc.styles['Заголовок1']
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        srf(run, 12, bold=True)

# P[4]: (ГБПОУ РО «СИТ») split into 3 runs
p = doc.paragraphs[4]
p.style = doc.styles['Заголовок1']
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.clear()
for part in ['(ГБПОУ ', 'РО «СИТ', '»)']:
    srf(p.add_run(part), 12, bold=True)

# КТП: 12pt (not 14pt!)
idx, p = find_para(doc, 'КАЛЕНДАРНО-ТЕМАТИЧЕСКИЙ ПЛАН')
if p:
    p.style = doc.styles['Заголовок1']
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    srf(p.add_run('КАЛЕНДАРНО-ТЕМАТИЧЕСКИЙ ПЛАН'), 12, bold=True)

# Semester/course line with underlines
idx, p = find_para(doc, 'семестр')
if p and 'курс' in p.text.lower():
    p.style = doc.styles['Заголовок1']
    p.alignment = None
    p.clear()
    def a(text, ul=False):
        srf(p.add_run(text), 11, underline=ul)
    a('на __'); a('3-4', ul=True); a('__ семест'); a('р('); a('ы) _')
    a('202', ul=True); a('6', ul=True); a('___ - __')
    a('202', ul=True); a('7', ul=True); a('___ учебного года _____')
    a('2', ul=True); a('______ курс')
    print('  Fixed semester line')

# Group line with underline for М-21
idx, p = find_para(doc, 'учебной группы')
if p:
    p.style = doc.styles['Заголовок1']
    p.alignment = None
    p.clear()
    def a(text, ul=False):
        srf(p.add_run(text), 11, underline=ul)
    a('учебной группы (учебных групп) ___________')
    a(' ', ul=True); a('М', ul=True); a('-', ul=True); a('2', ul=True); a('1', ul=True)
    a('____________________')
    print('  Fixed group line')

# ПМ line
idx, p = find_para(doc, 'Профессиональный модуль:')
if p:
    p.style = doc.styles['Normal']
    p.alignment = None
    p.clear()
    srf(p.add_run('Профессиональный модуль: '), 11)
    srf(p.add_run('ПМ.05 Выполнение работ по профессии 19861 Электромонтер по ремонту и обслуживанию электрооборудования промышленных и гражданских зданий и сооружений'), 11, underline=True)

# МДК line
idx, p = find_para(doc, 'Междисциплинарные курсы')
if p:
    p.style = doc.styles['Normal']
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.clear()
    srf(p.add_run('Междисциплинарные курсы: '), 11)
    srf(p.add_run('МДК 05.02 Организация и выполнение работ по сборке и монтажу электрооборудования по профессии 19861 Электромонтер по ремонту и обслуживанию электрооборудования промышленных и гражданских зданий и сооружений'), 11, underline=True)

# Specialty line
idx, p = find_para(doc, 'по профессии:')
if p and '08.02.09' in p.text:
    p.style = doc.styles['Normal']
    p.alignment = None
    p.clear()
    srf(p.add_run('по профессии: '), 11)
    srf(p.add_run('08.02.09 Монтаж, наладка и эксплуатация электрооборудования промышленных и гражданских зданий и сооружений'), 11, underline=True)

print('Changes 1-2: Title + Ministry DONE')

# ============================================================
# ИЗМЕНЕНИЕ 3: Таблица 1 - 4-row header, 12 rows
# ============================================================
body = doc.element.body

# Remove old Table 1 (tables[1]) and save its position reference
old_t1 = doc.tables[1]
old_t1_elem = old_t1._tbl

# Create new table element
from docx.oxml.table import CT_Tbl
new_tbl = doc.add_table(rows=12, cols=12, style='Table Grid')
new_t1_elem = new_tbl._tbl

# Move new table to where old T1 was
old_t1_elem.addnext(new_t1_elem)
body.remove(old_t1_elem)

# Now populate new Table 1
# Row 0
sc(new_tbl.rows[0], 0, 'Междисциплинарный курс (индекс МДК)')
sc(new_tbl.rows[0], 1, 'Курс')
sc(new_tbl.rows[0], 2, 'Семестр')
new_tbl.rows[0].cells[3].merge(new_tbl.rows[0].cells[9])
sc(new_tbl.rows[0], 3, 'Объём времени, отведённый на освоение междисциплинарного курса (по ФГОС)')
new_tbl.rows[0].cells[10].merge(new_tbl.rows[0].cells[11])
sc(new_tbl.rows[0], 10, 'Практика')

# Row 1
sc(new_tbl.rows[1], 0, 'Междисциплинарный курс (индекс МДК)')
sc(new_tbl.rows[1], 1, 'Курс'); sc(new_tbl.rows[1], 2, 'Семестр')
sc(new_tbl.rows[1], 3, 'Объем образовательной программы')
new_tbl.rows[1].cells[4].merge(new_tbl.rows[1].cells[8])
sc(new_tbl.rows[1], 4, 'Учебная нагрузка во взаимодействии с преподавателем')
sc(new_tbl.rows[1], 9, 'Самостоятельная работа обучающегося, часов')
sc(new_tbl.rows[1], 10, 'Учебная'); sc(new_tbl.rows[1], 11, 'Производственная')

# Row 2
for ci in [0,1,2]: sc(new_tbl.rows[2], ci, 'Междисциплинарный курс (индекс МДК)' if ci==0 else ['Курс','Семестр'][ci-1])
sc(new_tbl.rows[2], 0, 'Междисциплинарный курс (индекс МДК)')
sc(new_tbl.rows[2], 1, 'Курс'); sc(new_tbl.rows[2], 2, 'Семестр')
sc(new_tbl.rows[2], 3, 'Объем образовательной программы')
sc(new_tbl.rows[2], 4, 'Всего, часов')
new_tbl.rows[2].cells[5].merge(new_tbl.rows[2].cells[8])
sc(new_tbl.rows[2], 5, 'в т.ч.')
sc(new_tbl.rows[2], 9, 'Самостоятельная работа обучающегося, часов')
sc(new_tbl.rows[2], 10, 'Учебная'); sc(new_tbl.rows[2], 11, 'Производственная')

# Row 3 (detailed)
sc(new_tbl.rows[3], 0, 'Междисциплинарный курс (индекс МДК)')
sc(new_tbl.rows[3], 1, 'Курс'); sc(new_tbl.rows[3], 2, 'Семестр')
sc(new_tbl.rows[3], 3, 'Объем образовательной программы')
sc(new_tbl.rows[3], 4, 'Всего, часов')
sc(new_tbl.rows[3], 5, 'Теоретические занятия')
sc(new_tbl.rows[3], 6, 'лабораторные работы, часов')
sc(new_tbl.rows[3], 7, 'практические занятия, часов')
sc(new_tbl.rows[3], 8, 'Курсовые работы (проекты), часов')
sc(new_tbl.rows[3], 9, 'Самостоятельная работа обучающегося, часов')
sc(new_tbl.rows[3], 10, 'Учебная'); sc(new_tbl.rows[3], 11, 'Производственная')

# Row 4: column numbers
for ci in range(12):
    sc(new_tbl.rows[4], ci, str(ci + 1))

# Row 5-6: data
for ri, data in enumerate([
    ['МДК 05.02', '2', '3', '36', '36', '16', '-', '20', '-', '-', '-', '-'],
    ['МДК 05.02', '2', '4', '76', '76', '36', '-', '40', '-', '-', '-', '-'],
], 5):
    for ci, txt in enumerate(data):
        sc(new_tbl.rows[ri], ci, txt, bold=False)

# Rows 7-9: empty
for ri in range(7, 10):
    for ci in range(12):
        new_tbl.rows[ri].cells[ci].text = ''

# Row 10: Практика
for ci, txt in enumerate(['Практика', '-', '-', '-', '-', '-', '-', '-', '-', '-', '-', '-']):
    sc(new_tbl.rows[10], ci, txt, bold=(ci==0))

# Row 11: Всего
for ci, txt in enumerate(['Всего', '', '', '112', '112', '52', '-', '60', '-', '-', '-', '-']):
    sc(new_tbl.rows[11], ci, txt, bold=(ci==0))

print('Change 3: Table 1 rebuilt (4-row header, 12 rows) at correct position')

# ============================================================
# ИЗМЕНЕНИЕ 4: «Бардаков» в Table 2
# ============================================================
t2 = None
for t in doc.tables:
    if len(t.rows) > 50 and 'Наименование разделов' in t.rows[0].cells[1].text:
        t2 = t
        break

for ri, row in enumerate(t2.rows):
    if 'МДК 05.02' in row.cells[1].text and 'Организация и выполнение' in row.cells[1].text:
        cell = row.cells[-1]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run('Бардаков')
        r.font.size = Pt(10); r.font.name = 'Times New Roman'
        print(f'  Set Бардаков in T2 row {ri}')
        break

print('Change 4: Бардakov DONE')

# ============================================================
# ИЗМЕНЕНИЕ 5: «Защита пр.работы» → «Защита пр.раб.»
# ============================================================
count = 0
for row in t2.rows:
    cell = row.cells[9]
    if 'Защита пр.работы' in cell.text:
        for p in cell.paragraphs:
            for run in p.runs:
                if 'Защита пр.работы' in run.text:
                    run.text = run.text.replace('Защита пр.работы', 'Защита пр.раб.')
                    count += 1
print(f'Change 5: {count} replacements Защита пр.работы → Защита пр.раб.')

# ============================================================
# ИЗМЕНЕНИЕ 6: Одна строка диф.зачёта (no competencies)
# ============================================================
diff_header = None
for ri, row in enumerate(t2.rows):
    txt = row.cells[1].text
    if 'Дифференцированный зачет' in txt and row.cells[0].text.strip() == '':
        diff_header = ri
        break

if diff_header is not None:
    t2.rows[diff_header]._tr.getparent().remove(t2.rows[diff_header]._tr)
    print(f'  Removed empty диф.зачёт row [{diff_header}]')

for ri, row in enumerate(t2.rows):
    if 'Дифференцированный зачет' in row.cells[1].text and row.cells[0].text.strip():
        for ci in [4, 5, 6, 7, 8]:
            row.cells[ci].text = ''
        print(f'  Cleared competencies from диф.зачёт row [{ri}]')
        break

print('Change 6: Single диф.зачёт row DONE')

# ============================================================
# ИЗМЕНЕНИЕ 7: ДИ (3 книги) + ЭИ (5 ресурсов с URL)
# ============================================================
# Remove old mixed table (last table)
old_lit = doc.tables[-1]
print(f'  Removing old literature table ({len(old_lit.rows)} rows)')
old_lit_pos = old_lit._tbl
old_lit_pos.addnext  # just to verify it exists

# We need to insert new tables WHERE the old one was
# First remove old, then add new at same position
old_lit_elem = old_lit._tbl
ref_element = old_lit_elem.getnext()  # element after old table (or None)
parent = old_lit_elem.getparent()
parent.remove(old_lit_elem)

# Create ДИ table (4 cols, 4 rows)
di_data = [
    ('ДИ 1', 'Общая технология электромонтажных работ: учебник для СПО', 'Григорьева С.В.', 'М.: ИЦ "Академия", 2020'),
    ('ДИ 2', 'Организация и выполнение работ по монтажу и наладке электрооборудования: учебное пособие', 'Бычков А.В.', 'М.: ИЦ "Академия", 2020'),
    ('ДИ 3', 'Организация и выполнение работ по монтажу и наладке электрооборудования: учебное пособие', 'Шашкова И.В., Бычков А.В.', 'М.: ИЦ "Академия", 2020'),
]
di_tbl = doc.add_table(rows=4, cols=4, style='Table Grid')
for ci, h in enumerate(['№ п/п', 'Наименование', 'Автор', 'Издательство, год издания']):
    cell = di_tbl.rows[0].cells[ci]; cell.text = ''
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h); r.font.size = Pt(10); r.font.name = 'Times New Roman'; r.bold = True

for ri, (num, name, author, year) in enumerate(di_data, 1):
    sc(di_tbl.rows[ri], 0, num, bold=False)
    sc(di_tbl.rows[ri], 1, name, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    sc(di_tbl.rows[ri], 2, author, bold=False)
    sc(di_tbl.rows[ri], 3, year, bold=False)

# Create ЭИ table (5 cols, 6 rows)
ei_data = [
    ('ЭИ 1', 'Электрические системы и сети. Энергосбережение: учебник', 'Климова Г.Н.', 'М.: Издательство Юрайт, 2023', 'https://urait.ru/bcode/513864'),
    ('ЭИ 2', 'Организация и методика производственного обучения.', 'Бредихин А.Н.', 'М.: Издательство Юрайт, 2023', 'https://urait.ru/bcode/517783'),
    ('ЭИ 3', 'Информационный портал для электромонтеров', '', '', 'http://electromonter.info'),
    ('ЭИ 4', 'Образовательный сайт "Школа для электрика"', '', '', 'http://ElectricalSchool.info'),
    ('ЭИ 5', 'Нормативно-технические документы', '', '', 'http://electrolibrary.info'),
]
ei_tbl = doc.add_table(rows=6, cols=5, style='Table Grid')
for ci, h in enumerate(['№ п/п', 'Наименование', 'Автор', 'Издательство, год издания', 'URL (ссылка)']):
    cell = ei_tbl.rows[0].cells[ci]; cell.text = ''
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h); r.font.size = Pt(10); r.font.name = 'Times New Roman'; r.bold = True

for ri, (num, name, author, year, url) in enumerate(ei_data, 1):
    sc(ei_tbl.rows[ri], 0, num, bold=False)
    sc(ei_tbl.rows[ri], 1, name, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT)
    sc(ei_tbl.rows[ri], 2, author, bold=False)
    sc(ei_tbl.rows[ri], 3, year, bold=False)
    sc(ei_tbl.rows[ri], 4, url, bold=False)

# Move ДИ and ЭИ tables to where old table was
# add_table appends to end, so we need to move them
if ref_element is not None:
    ref_element.addprevious(di_tbl._tbl)
    di_tbl._tbl.addnext(ei_tbl._tbl)
else:
    # No element after old table, just append (they're already at end)
    pass

print('Change 7: ДИ (3) + ЭИ (5 with URL) DONE')

# ============================================================
# SAVE
# ============================================================
doc.save(DOC_PATH)
print(f'\nAll 7 changes applied → {DOC_PATH}')
