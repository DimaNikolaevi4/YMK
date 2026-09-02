from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json

DST = '/home/z/my-project/YMK/KTP/МДК_05_02_КТП.docx'

# ========== LOAD DATA ==========
with open('/tmp/table2_data.json', 'r', encoding='utf-8') as f:
    TABLE2_DATA = json.load(f)

FONT = 'Times New Roman'


def set_cell_margins(cell, top=0, bottom=0, left=40, right=40):
    """Set cell margins using OxmlElement (minimal XML, safe)."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is not None:
        tcPr.remove(tcMar)
    el = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        m = OxmlElement('w:' + side)
        m.set(qn('w:w'), str(val))
        m.set(qn('w:type'), 'dxa')
        el.append(m)
    tcPr.append(el)


def set_cell_vertical_align(cell, align='center'):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    va = tcPr.find(qn('w:vAlign'))
    if va is not None:
        tcPr.remove(va)
    v = OxmlElement('w:vAlign')
    v.set(qn('w:val'), align)
    tcPr.append(v)


def set_cell_border(cell, **kwargs):
    """Set cell borders. Usage: set_cell_border(cell, top={...}, ...)"""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    borders = tcPr.find(qn('w:tcBorders'))
    if borders is not None:
        tcPr.remove(borders)
    borders = OxmlElement('w:tcBorders')
    for edge, attrs in kwargs.items():
        el = OxmlElement('w:' + edge)
        for k, v in attrs.items():
            el.set(qn('w:' + k), str(v))
        borders.append(el)
    tcPr.append(borders)


def set_table_borders(table):
    """Set all borders for a table."""
    border_def = {'val': 'single', 'sz': '4', 'space': '0', 'color': '000000'}
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    borders = OxmlElement('w:tblBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement('w:' + edge)
        for k, v in border_def.items():
            el.set(qn('w:' + k), v)
        borders.append(el)
    tblPr.append(borders)


def set_cell_width(cell, width_cm):
    """Set column width in cm."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    tw = tcPr.find(qn('w:tcW'))
    if tw is not None:
        tcPr.remove(tw)
    el = OxmlElement('w:tcW')
    el.set(qn('w:w'), str(int(width_cm * 567)))  # cm -> twips
    el.set(qn('w:type'), 'dxa')
    tcPr.append(el)


def fmt_cell(cell, text, size=10, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    """Format a table cell with text."""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(14)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold


def fmt_cell_multiline(cell, lines, size=10, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    """Format a cell with multiple lines (separated by line breaks)."""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(14)
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = FONT
        run.font.size = Pt(size)
        run.font.bold = bold


def fmt_all_cells(table, rows_range=None, size=10):
    """Apply font formatting to all cells in a table."""
    for ri, row in enumerate(table.rows):
        if rows_range and ri not in rows_range:
            continue
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    if r.font.name is None:
                        r.font.name = FONT
                    if r.font.size is None:
                        r.font.size = Pt(size)


def add_formatted_paragraph(doc, text, size=11, bold=False, alignment=None):
    """Add a paragraph with specific formatting."""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    return p


def main():
    doc = Document()

    # ========== PAGE SETUP ==========
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(1.5)

    # ========== DEFAULT STYLE ==========
    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(11)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    # ========== TITLE PAGE ==========
    add_formatted_paragraph(doc, 'МИНИСТЕРСТВО ОБЩЕГО И ПРОФЕССИОНАЛЬНОГО ОБРАЗОВАНИЯ РОСТОВСКОЙ ОБЛАСТИ', size=11, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_formatted_paragraph(doc, 'ГОСУДАРСТВЕННОЕ БЮДЖЕТНОЕ ПРОФЕССИОНАЛЬНОЕ ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ РОСТОВСКОЙ ОБЛАСТИ', size=11, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_formatted_paragraph(doc, '«САЛЬСКИЙ ИНДУСТРИАЛЬНЫЙ ТЕХНИКУМ»', size=11, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    add_formatted_paragraph(doc, '(ГБПОУ РО «СИТ»)', size=11, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    doc.add_paragraph()

    # Table 0: УТВЕРЖДАЮ (3x3)
    t0 = doc.add_table(rows=3, cols=3)
    t0.alignment = WD_TABLE_ALIGNMENT.RIGHT
    fmt_cell(t0.cell(1, 1), 'УТВЕРЖДАЮ', size=11, bold=True)
    fmt_cell(t0.cell(1, 2), '____________ Т.В. Якимова', size=11)
    fmt_cell(t0.cell(2, 2), '«____» ______________ 20 ____ г.', size=11)
    # Remove borders from table 0
    for row in t0.rows:
        for cell in row.cells:
            set_cell_border(cell, top={'val': 'none', 'sz': '0', 'color': 'auto'},
                                  bottom={'val': 'none', 'sz': '0', 'color': 'auto'},
                                  left={'val': 'none', 'sz': '0', 'color': 'auto'},
                                  right={'val': 'none', 'sz': '0', 'color': 'auto'})

    doc.add_paragraph()
    add_formatted_paragraph(doc, 'КАЛЕНДАРНО-ТЕМАТИЧЕСКИЙ ПЛАН', size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()

    title_lines = [
        'на 3, 4 семестры 2026-2027 учебного года 2 курс',
        '',
        'учебной группы (учебных групп) М-21',
        '',
        'Профессиональный модуль: ПМ.05 Выполнение работ по профессии 19861 Электромонтер по ремонту и обслуживанию электрооборудования',
        '',
        'Междисциплинарные курсы: МДК 05.02 Организация и выполнение работ по сборке и монтажу электрооборудования и распределительных устройств',
        '',
        'по профессии: 08.02.09 Монтаж, наладка и эксплуатация электрооборудования промышленных и гражданских зданий',
        '',
        'Объем образовательной программы: ______112______________ (часов);',
        'в том числе в форме практической подготовки ___60______________(часа);',
        '',
        'Учебная нагрузка во взаимодействии с преподавателем ________112_______________ (часов):',
        'из нее:',
        'теоретическое обучение __52___ (часов);                    практические занятия __60__ (часов);',
        'лабораторные занятия ____________ (часов);                   курсовая работа/проект ____________ (часов);',
        'самостоятельная работа ____________ (часов);',
        '',
        'промежуточная аттестация в форме ___Комплексный дифференцированный зачет___',
        '                                                                               (указать форму)',
        '',
        'Составлен в соответствии с рабочей программой ПМ 05, утверждённой __________________',
        '              (дата утверждения)',
        '',
        'Рассмотрен на заседании цикловой комиссии _____________________________ дисциплин',
        'Протокол  №_______от________________20_____ года',
        'Председатель цикловой комиссии ___________________________/Ткаченко А.Н./',
        '',
        '',
    ]
    for line in title_lines:
        if line == '':
            doc.add_paragraph()
        else:
            add_formatted_paragraph(doc, line, size=11)

    # г. Сальск and year - right aligned
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('г. Сальск')
    run.font.name = FONT
    run.font.size = Pt(11)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('2026-2027 уч. год')
    run.font.name = FONT
    run.font.size = Pt(11)

    doc.add_paragraph()

    # ========== TABLE 1: Распределение часов ==========
    add_formatted_paragraph(doc, 'Распределение часов по профессиональному модулю', size=11)
    add_formatted_paragraph(doc, 'Таблица 1', size=11)
    doc.add_paragraph()

    t1_headers = ['Междисциплинарный курс (индекс МДК)', 'Курс', 'Семестр',
                   'Объём образовательной программы', 'Всего, часов', 'Теоретические занятия',
                   'Лабораторные работы, часов', 'Практические занятия, часов',
                   'Курсовые работы (проекты), часов', 'Самостоятельная работа обучающегося, часов',
                   'Учебная практика, часов', 'Производственная практика, часов']
    t1_data = [
        ['МДК 05.02', '2', '3', '36', '36', '16', '-', '20', '-', '-', '-', '-'],
        ['МДК 05.02', '2', '4', '76', '76', '36', '-', '40', '-', '-', '-', '-'],
        ['Практика', '-', '-', '-', '-', '-', '-', '-', '-', '-', '-', '-'],
        ['Всего', '', '', '112', '112', '52', '-', '60', '-', '-', '-', '-'],
    ]

    t1 = doc.add_table(rows=1 + len(t1_data), cols=len(t1_headers))
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t1)

    for ci, h in enumerate(t1_headers):
        fmt_cell(t1.cell(0, ci), h, size=9, bold=True)
    for ri, row_data in enumerate(t1_data):
        for ci, val in enumerate(row_data):
            fmt_cell(t1.cell(ri + 1, ci), val, size=9)
    fmt_all_cells(t1, range(0, len(t1_data) + 1), size=9)

    doc.add_paragraph()
    add_formatted_paragraph(doc, 'Форма промежуточной аттестации обучающихся за семестр по междисциплинарному курсу', size=11)
    add_formatted_paragraph(doc, '(МДК 05.02 Организация и выполнение работ по сборке и монтажу электрооборудования и распределительных устройств) – (Комплексный дифференцированный зачет).', size=11)
    doc.add_paragraph()

    # ========== TABLE 2: Содержание обучения (11 columns) ==========
    add_formatted_paragraph(doc, 'Содержание обучения по профессиональному модулю', size=11)
    # "Таблица 2" right-aligned
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('Таблица 2')
    run.font.name = FONT
    run.font.size = Pt(11)
    doc.add_paragraph()

    T2_COLS = 11
    T2_HEADERS_TEXT = [
        ['№ занятия', 'Наименование разделов\nпрофессионального модуля,\nтем и занятий по МДК', 'Обязательная учебная нагрузка', '', '', 'Коды формируемых компетенции', '', 'Материальное и информационное\nобеспечение занятий', 'Задания для студентов', 'Формы и методы контроля', 'ФИО преподавателя'],
        ['', '', 'Кол-во\nчасов', 'В форме практической подготовки', 'Вид занятия', '', '', '', '', '', ''],
        ['', '', '', '', '', 'ОК', 'ПК', '', '', '', ''],
        ['1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11'],
    ]

    # Skip first 2 rows of old data (old header rows)
    table2_data = TABLE2_DATA[2:]  # Skip old number row and header row
    # Calculate rows needed: 4 header + data rows
    num_data_rows = len(table2_data)
    total_rows = 4 + num_data_rows

    t2 = doc.add_table(rows=total_rows, cols=T2_COLS)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t2)

    # ---- Fill header rows ----
    # Row 0
    header_row0_texts = ['№ занятия', None, 'Обязательная учебная нагрузка', None, None,
                         'Коды формируемых компетенции', None,
                         'Материальное и информационное обеспечение занятий',
                         'Задания для студентов', 'Формы и методы контроля', 'ФИО преподавателя']

    # Col 0: № занятия
    fmt_cell(t2.cell(0, 0), '№ занятия', size=10, bold=True)
    # Col 1: multiline
    fmt_cell_multiline(t2.cell(0, 1),
                       ['Наименование разделов', 'профессионального модуля,', 'тем и занятий по МДК'],
                       size=10, bold=True)
    # Col 2: Обязательная учебная нагрузка (will h-merge with 3,4)
    fmt_cell(t2.cell(0, 2), 'Обязательная учебная нагрузка', size=10, bold=True)
    # Cols 3,4 will be hidden by merge
    # Col 5: Коды формируемых компетенции (will h-merge with 6)
    fmt_cell(t2.cell(0, 5), 'Коды формируемых компетенции', size=10, bold=True)
    # Col 7: multiline
    fmt_cell_multiline(t2.cell(0, 7),
                       ['Материальное и информационное', 'обеспечение занятий'],
                       size=10, bold=True)
    fmt_cell(t2.cell(0, 8), 'Задания для студентов', size=10, bold=True)
    fmt_cell(t2.cell(0, 9), 'Формы и методы контроля', size=10, bold=True)
    fmt_cell(t2.cell(0, 10), 'ФИО преподавателя', size=10, bold=True)

    # Row 1
    fmt_cell_multiline(t2.cell(1, 2), ['Кол-во', 'часов'], size=10, bold=True)
    fmt_cell(t2.cell(1, 3), 'В форме практической подготовки', size=10, bold=True)
    fmt_cell(t2.cell(1, 4), 'Вид занятия', size=10, bold=True)

    # Row 2
    fmt_cell(t2.cell(2, 5), 'ОК', size=10, bold=True)
    fmt_cell(t2.cell(2, 6), 'ПК', size=10, bold=True)

    # Row 3: numbers
    for ci, num in enumerate(['1','2','3','4','5','6','7','8','9','10','11']):
        fmt_cell(t2.cell(3, ci), num, size=10, bold=True)

    # ---- Apply MERGES (header) ----
    # Vertical merges: col 0,1,7,8,9,10 rows 0-2
    t2.cell(0, 0).merge(t2.cell(2, 0))
    t2.cell(0, 1).merge(t2.cell(2, 1))
    t2.cell(0, 7).merge(t2.cell(2, 7))
    t2.cell(0, 8).merge(t2.cell(2, 8))
    t2.cell(0, 9).merge(t2.cell(2, 9))
    t2.cell(0, 10).merge(t2.cell(2, 10))

    # Horizontal merge: row 0, cols 2-4 (Обязательная учебная нагрузка)
    t2.cell(0, 2).merge(t2.cell(0, 4))

    # Horizontal merge: cols 5-6 in rows 0 and 1
    t2.cell(0, 5).merge(t2.cell(0, 6))
    t2.cell(1, 5).merge(t2.cell(1, 6))
    # Vertical merge the h-merged cells (0,5)-(0,6) with (1,5)-(1,6)
    # After h-merge, (0,5) spans cols 5-6; same for (1,5)
    t2.cell(0, 5).merge(t2.cell(1, 5))

    # Vertical merges for col 2,3,4 rows 1-2
    t2.cell(1, 2).merge(t2.cell(2, 2))
    t2.cell(1, 3).merge(t2.cell(2, 3))
    t2.cell(1, 4).merge(t2.cell(2, 4))

    # ---- Fill data rows ----
    for i, data in enumerate(table2_data):
        ri = 4 + i
        num_pp = data[0]
        name = data[1]
        hours = data[2]
        pract = data[3]
        vid = data[4]
        osnash = data[5]
        zadanie = data[6]
        prim = data[7]

        is_header = (num_pp == '' and vid == '')
        is_total = name.startswith('Итого')
        bold = is_header or is_total
        name_align = WD_ALIGN_PARAGRAPH.LEFT

        fmt_cell(t2.cell(ri, 0), num_pp, size=10, bold=bold)
        fmt_cell(t2.cell(ri, 1), name, size=10, bold=bold, align=name_align)
        fmt_cell(t2.cell(ri, 2), hours, size=10, bold=bold)
        fmt_cell(t2.cell(ri, 3), pract, size=10, bold=bold)
        fmt_cell(t2.cell(ri, 4), vid, size=10, bold=bold)
        fmt_cell(t2.cell(ri, 5), '', size=10)  # ОК
        fmt_cell(t2.cell(ri, 6), '', size=10)  # ПК
        fmt_cell(t2.cell(ri, 7), osnash, size=10)  # Мат. обеспечение
        fmt_cell(t2.cell(ri, 8), zadanie, size=10)  # Задания
        fmt_cell(t2.cell(ri, 9), prim, size=10)  # Формы контроля
        fmt_cell(t2.cell(ri, 10), '', size=10)  # ФИО преподавателя

    # Apply formatting to all cells in table 2
    for ri in range(total_rows):
        for ci in range(T2_COLS):
            try:
                cell = t2.cell(ri, ci)
                set_cell_vertical_align(cell, 'center')
            except Exception:
                pass

    fmt_all_cells(t2, range(total_rows), size=10)

    doc.add_paragraph()

    # ========== TABLE 3 (2а): Материально-техническое обеспечение ==========
    add_formatted_paragraph(doc, 'Материально-техническое обеспечение занятий', size=11)
    add_formatted_paragraph(doc, 'Таблица 2а', size=11)
    doc.add_paragraph()

    t3_data = [
        ['Кабинет физики, электротехники и электроники'],
        ['Лаборатория электротехники и электроники'],
        ['Методические указания по выполнению лабораторно-практических работ'],
        ['Лаборатория электрических измерений и электрических цепей'],
        ['Мастерская \u00abСлесарно-механическая\u00bb'],
        ['Мастерская \u00abЭлектротехническая\u00bb'],
        ['Мастерская \u00abМонтажа, технического обслуживания и эксплуатации электрооборудования\u00bb'],
        ['Средства мультимедиа'],
        ['ЭВМ'],
    ]

    t3 = doc.add_table(rows=1 + len(t3_data), cols=2)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t3)
    fmt_cell(t3.cell(0, 0), '№ п/п', size=10, bold=True)
    fmt_cell(t3.cell(0, 1), 'Материально-техническое обеспечение занятий', size=10, bold=True)
    for ri, row_d in enumerate(t3_data):
        fmt_cell(t3.cell(ri + 1, 0), str(ri + 1), size=10)
        fmt_cell(t3.cell(ri + 1, 1), row_d[0], size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    fmt_all_cells(t3, range(len(t3_data) + 1), size=10)

    doc.add_paragraph()

    # ========== TABLE 4 (2б): Основные источники ==========
    add_formatted_paragraph(doc, 'Информационное обеспечение обучения', size=11)
    add_formatted_paragraph(doc, 'Основные источники (ОИ):', size=11)
    add_formatted_paragraph(doc, 'Таблица 2б', size=11)
    doc.add_paragraph()

    t4_data = [
        ['ОИ 1', 'Общая технология электромонтажных работ: учебник для СПО', 'Григорьева С.В.', 'М.: ИЦ \u00abАкадемия\u00bb, 2020'],
        ['ОИ 2', 'Технология электромонтажных работ: учеб. пособие для СПО', 'Нестеренко В.М.', 'М.: ИЦ \u00abАкадемия\u00bb, 2022'],
        ['ОИ 3', 'Сборка, монтаж, регулировка и ремонт узлов и механизмов оборудования, агрегатов, машин, станков и другого электрооборудования промышленных организаций: учебник', 'Сидорова Л.Г.', 'М.: ИЦ \u00abАкадемия\u00bb, 2022'],
        ['ОИ 4', 'Проверка и наладка электрооборудования: учебник', 'Ярочкина Г.В.', 'М.: ИЦ \u00abАкадемия\u00bb, 2022'],
        ['ОИ 5', 'Организация и выполнение работ по монтажу и наладке электрооборудования промышленных и гражданских зданий. В двух частях. Часть 1. Внутреннее электроснабжение промышленных и гражданских зданий: учебник', 'Бычков А.В.', 'М.: ИЦ \u00abАкадемия\u00bb, 2020'],
        ['ОИ 6', 'Организация и выполнение работ по монтажу и наладке электрооборудования промышленных и гражданских зданий. В двух частях. Часть 2. Монтаж и наладка электрооборудования промышленных и гражданских зданий: учебник', 'Шашкова И.В., Бычков А.В.', 'М.: ИЦ \u00abАкадемия\u00bb, 2020'],
        ['ОИ 7', 'Монтаж, наладка, эксплуатация и ремонт систем электроснабжения промышленных предприятий: учебное пособие для СПО', 'Полуянович Н.К.', 'Санкт-Петербург: Лань, 2022'],
    ]

    t4 = doc.add_table(rows=1 + len(t4_data), cols=4)
    t4.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t4)
    fmt_cell(t4.cell(0, 0), '№ п/п', size=10, bold=True)
    fmt_cell(t4.cell(0, 1), 'Наименование', size=10, bold=True)
    fmt_cell(t4.cell(0, 2), 'Автор', size=10, bold=True)
    fmt_cell(t4.cell(0, 3), 'Издательство, год издания', size=10, bold=True)
    for ri, row_d in enumerate(t4_data):
        for ci, val in enumerate(row_d):
            al = WD_ALIGN_PARAGRAPH.LEFT if ci >= 1 else WD_ALIGN_PARAGRAPH.CENTER
            fmt_cell(t4.cell(ri + 1, ci), val, size=9, align=al)
    fmt_all_cells(t4, range(len(t4_data) + 1), size=9)

    doc.add_paragraph()

    # ========== TABLE 5 (2в): Дополнительные источники ==========
    add_formatted_paragraph(doc, 'Дополнительные источники (ДИ):', size=11)
    add_formatted_paragraph(doc, 'Таблица 2в', size=11)
    doc.add_paragraph()

    t5_data = [
        ['ДИ 1', 'Электрические системы и сети. Энергосбережение: учебное пособие для СПО', 'Климова Г.Н.', 'М.: Издательство Юрайт, 2023'],
        ['ДИ 2', 'Организация и методика производственного обучения. Электромонтер-кабельщик: учебное пособие для СПО', 'Бредихин А.Н.', 'М.: Издательство Юрайт, 2023'],
        ['ДИ 3', 'Информационный портал для электромонтеров', '', 'http://electromonter.info'],
        ['ДИ 4', 'Образовательный сайт \u00abШкола для электрика\u00bb', '', 'http://ElectricalSchool.info'],
        ['ДИ 5', 'Нормативно-технические документы', '', 'http://electrolibrary.info'],
    ]

    t5 = doc.add_table(rows=1 + len(t5_data), cols=4)
    t5.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t5)
    fmt_cell(t5.cell(0, 0), '№ п/п', size=10, bold=True)
    fmt_cell(t5.cell(0, 1), 'Наименование', size=10, bold=True)
    fmt_cell(t5.cell(0, 2), 'Автор', size=10, bold=True)
    fmt_cell(t5.cell(0, 3), 'Издательство, год издания', size=10, bold=True)
    for ri, row_d in enumerate(t5_data):
        for ci, val in enumerate(row_d):
            al = WD_ALIGN_PARAGRAPH.LEFT if ci >= 1 else WD_ALIGN_PARAGRAPH.CENTER
            fmt_cell(t5.cell(ri + 1, ci), val, size=9, align=al)
    fmt_all_cells(t5, range(len(t5_data) + 1), size=9)

    # ========== SAVE ==========
    doc.save(DST)

    # ========== VERIFY ==========
    import zipfile
    import os
    fsize = os.path.getsize(DST)
    print('File size:', fsize, 'bytes ({:.1f} KB)'.format(fsize / 1024))
    with zipfile.ZipFile(DST) as z:
        names = z.namelist()
        print('ZIP entries:', len(names))
        for n in sorted(names):
            print('  ' + n + ' (' + str(z.getinfo(n).file_size) + ')')

    # Verify tables
    doc2 = Document(DST)
    print('Tables:', len(doc2.tables))
    for ti, t in enumerate(doc2.tables):
        print('  T' + str(ti) + ': ' + str(len(t.rows)) + 'r x ' + str(len(t.columns)) + 'c')
    print('OK')


if __name__ == '__main__':
    main()
