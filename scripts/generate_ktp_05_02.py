#!/usr/bin/env python3
"""
Генерация КТП МДК 05.02 на семестры 3+4 для группы М-21.
Специальность 08.02.09 Монтаж, наладка и эксплуатация электрооборудования.

Принципы создания КТП:
- Таблица 2: 11 колонок (по эталону КТП МДК 02.01)
- Столбец 8 (Мат.обеспечение): для теории — ОИ2 (монтаж 2.2-2.4) или ОИ3 (безопасность/ремонт 2.1, 2.5-2.8),
  для практики — кабинет/лаборатория (№3, 2а)
- Столбец 9 (Задания): для теории — страницы из ОИ2/ОИ3, оценочные (±3-5 стр.);
  для практики — методические указания
- Заголовок Таблицы 2: 3 строки с объединениями (gridSpan/vMerge через
  python-docx .merge())
"""
import sys
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# УТИЛИТЫ
# ============================================================

def set_cell_shading(cell, color):
    """Заливка ячейки цветом."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_text(cell, text, bold=False, size=8, alignment=WD_ALIGN_PARAGRAPH.LEFT, font_name='Times New Roman'):
    """Установить текст ячейки с форматированием."""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(size + 3)
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    run.font.name = font_name
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)


def add_paragraph(doc, text, bold=False, size=11, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')
    return p


def set_col_width(cell, width_emu):
    """Установить ширину столбца."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(width_emu))
    tcW.set(qn('w:type'), 'dxa')
    # Remove existing tcW if any
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    tcPr.append(tcW)


# ============================================================
# ДАННЫЕ МДК 05.02 — извлечены из РП ПМ.05
# ============================================================

MDK_DATA = {
    "mdk_code": "05.02",
    "mdk_name": "Организация и выполнение работ по сборке и монтажу электрооборудования и распределительных устройств",
    "pm_code": "05",
    "pm_name": "Выполнение работ по профессии 19861 Электромонтер по ремонту и обслуживанию электрооборудования",
    "profession": "08.02.09 Монтаж, наладка и эксплуатация электрооборудования промышленных и гражданских зданий",
    "semester": "3, 4",
    "course": "2",
    "group": "М-21",
    "year": "2026-2027",
    "total_hours": 112,
    "practice_hours": 60,
    "with_teacher_hours": 112,
    "theory_hours": 52,
    "practical_hours": 60,
    "lab_hours": 0,
    "coursework_hours": 0,
    "independent_hours": 0,
    "exam_form": "Комплексный дифференцированный зачет",
    "commission_chair": "Ткаченко А.Н.",
    "director": "Т.В. Якимова",
    "semesters": [
        {"sem": 3, "total": 36, "theory": 16, "practice": 20},
        {"sem": 4, "total": 76, "theory": 36, "practice": 40},
    ],
}

# ============================================================
# ИСТОЧНИКИ ДЛЯ ТЕОРИИ (столбец 8)
# ============================================================
# ОИ2 (Нестеренко) — для монтажных тем: 2.2, 2.3, 2.4
# ОИ3 (Сидорова) — для безопасности (2.1) и ремонта (2.5–2.8)
# Страницы оценочные (±3–5 стр.), сверить с печатными экземплярами.
BOOK_OI2 = "ОИ2"
BOOK_OI3 = "ОИ3"
# Для практики — кабинет/лаборатория
PRACTICE_EQUIP = "№3, 2а"

# ============================================================
# ТЕМЫ И ЗАНЯТИЯ
# Сумма: 52 теория + 60 практика = 112
# Сем.3 = 36ч (Темы 2.1 + 2.2 + начало 2.3)
# Сем.4 = 76ч (окончание 2.3 + 2.4-2.8 + диф.зач)
#
# Столбец 8 (equipment):
#   теория -> BOOK_OI2 (темы 2.2-2.4 монтаж) или BOOK_OI3 (тема 2.1, 2.5-2.8 ремонт)
#   практика -> PRACTICE_EQUIP (№3, 2а — кабинет/лаборатория)
#
# Столбец 9 (task):
#   теория -> страницы из ОИ2 или ОИ3, согласованные с темой занятия (оценочные)
#   практика -> "Метод. указания"
# ============================================================

TOPICS = [
    # ========= Тема 2.1 (4ч: 2 теория + 2 практика) — СЕМ.3 =========
    {
        "name": "Тема 2.1 Безопасность труда при организации работ по сборке, монтажу и ремонту электрооборудования",
        "total_hours": 4,
        "practice_hours": 2,
        "semester": 3,
        "lessons": [
            {
                "name": "Тема 2.1 Безопасность труда при организации работ по сборке, монтажу и ремонту электрооборудования",
                "hours": 2, "is_practice": False, "is_lab": False,
                "practice_prep": 0,
                "equipment": BOOK_OI3,
                "task": "Стр. 5-89",
                "note": "Вх.контр"
            },
            {
                "name": "Практическое занятие №1 Правила оказания первой медицинской помощи при поражении электрическим током",
                "hours": 2, "is_practice": True, "is_lab": False,
                "practice_prep": 2,
                "equipment": PRACTICE_EQUIP,
                "task": "Метод. указания",
                "note": "Защита пр.работы"
            },
        ]
    },
    # ========= Тема 2.2 (30ч: 12 теория + 18 практика) — СЕМ.3 =========
    {
        "name": "Тема 2.2 Сборка и монтаж осветительных электроустановок и аппаратов защиты и пускорегулирующей аппаратуры",
        "total_hours": 30,
        "practice_hours": 18,
        "semester": 3,
        "lessons": [
            # Теория (12ч = 6 занятий по 2ч)
            {"name": "Тема 2.2 Сборка и монтаж осветительных электроустановок и аппаратов защиты и пускорегулирующей аппаратуры (часть 1)",
             "hours": 2, "is_practice": False, "is_lab": False, "practice_prep": 0,
             "equipment": BOOK_OI2, "task": "Стр. 79-93", "note": "Вх.контр"},
            {"name": "Тема 2.2 Сборка и монтаж осветительных электроустановок и аппаратов защиты и пускорегулирующей аппаратуры (часть 2)",
             "hours": 2, "is_practice": False, "is_lab": False, "practice_prep": 0,
             "equipment": BOOK_OI2, "task": "Стр. 94-108", "note": "Текущий контроль"},
            {"name": "Тема 2.2 Сборка и монтаж осветительных электроустановок и аппаратов защиты и пускорегулирующей аппаратуры (часть 3)",
             "hours": 2, "is_practice": False, "is_lab": False, "practice_prep": 0,
             "equipment": BOOK_OI2, "task": "Стр. 109-123", "note": "Текущий контроль"},
            {"name": "Тема 2.2 Сборка и монтаж осветительных электроустановок и аппаратов защиты и пускорегулирующей аппаратуры (часть 4)",
             "hours": 2, "is_practice": False, "is_lab": False, "practice_prep": 0,
             "equipment": BOOK_OI2, "task": "Стр. 124-138", "note": "Текущий контроль"},
            {"name": "Тема 2.2 Сборка и монтаж осветительных электроустановок и аппаратов защиты и пускорегулирующей аппаратуры (часть 5)",
             "hours": 2, "is_practice": False, "is_lab": False, "practice_prep": 0,
             "equipment": BOOK_OI2, "task": "Стр. 139-153", "note": "Текущий контроль"},
            {"name": "Тема 2.2 Сборка и монтаж осветительных электроустановок и аппаратов защиты и пускорегулирующей аппаратуры (часть 6)",
             "hours": 2, "is_practice": False, "is_lab": False, "practice_prep": 0,
             "equipment": BOOK_OI2, "task": "Стр. 154-167", "note": "Текущий контроль"},
            # Практика (18ч = 9 занятий по 2ч)
            {"name": "Практическое занятие №2 Выполнение расчета выбора проводов осветительных установок",
             "hours": 2, "is_practice": True, "is_lab": False, "practice_prep": 2,
             "equipment": PRACTICE_EQUIP, "task": "Метод. указания", "note": "Защита пр.работы"},
            {"name": "Практическое занятие №3 Изучение технологии проверки исправности ламп и ПРА",
             "hours": 2, "is_practice": True, "is_lab": False, "practice_prep": 2,
             "equipment": PRACTICE_EQUIP, "task": "Метод. указания", "note": "Защита пр.работы"},
            {"name": "Практическое занятие №4 Сборка и проверка цепей электрического освещения",
             "hours": 2, "is_practice": True, "is_lab": False, "practice_prep": 2,
             "equipment": PRACTICE_EQUIP, "task": "Метод. указания", "note": "Защита пр.работы"},
            {"name": "Практическое занятие №5 Сборка схемы освещения",
             "hours": 2, "is_practice": True, "is_lab": False, "practice_prep": 2,
             "equipment": PRACTICE_EQUIP, "task": "Метод. указания", "note": "Защита пр.работы"},
            {"name": "Практическое занятие №6 Сборка и проверка цепей электрических распределительных щитов жилых и офисных помещений",
             "hours": 2, "is_practice": True, "is_lab": False, "practice_prep": 2,
             "equipment": PRACTICE_EQUIP, "task": "Метод. указания", "note": "Защита пр.работы"},
            {"name": "Практическое занятие №7 Изучение технологии монтажа и принципиальных схем включения осветительных электроустановок",
             "hours": 2, "is_practice": True, "is_lab": False, "practice_prep": 2,
             "equipment": PRACTICE_EQUIP, "task": "Метод. указания", "note": "Защита пр.работы"},
            {"name": "Практическое занятие №8 Выбор типа автоматического воздушного выключателя и тока его расцепителя",
             "hours": 2, "is_practice": True, "is_lab": False, "practice_prep": 2,
             "equipment": PRACTICE_EQUIP, "task": "Метод. указания", "note": "Защита пр.работы"},
            {"name": "Практическое занятие №9 Расчет плавкой вставки предохранителя и выбор типа предохранителя",
             "hours": 2, "is_practice": True, "is_lab": False, "practice_prep": 2,
             "equipment": PRACTICE_EQUIP, "task": "Метод. указания", "note": "Защита пр.работы"},
            {"name": "Практическое занятие №10 Электромонтаж и наладка магнитных пускателей",
             "hours": 2, "is_practice": True, "is_lab": False, "practice_prep": 2,
             "equipment": PRACTICE_EQUIP, "task": "Метод. указания", "note": "Защита пр.работы"},
        ]
    },
    # ========= Тема 2.3 (14ч: 8 теория + 6 практика) — СЕМ.3(2ч) + СЕМ.4(12ч) =========
    {
        "name": "Тема 2.3 Монтаж кабельных линий, комплектных шинопроводов и троллейных линий",
        "total_hours": 14,
        "practice_hours": 6,
        "semester": [3, 4],
        "lessons": [
            # Теория (8ч = 4 занятия)
            {"name": "Тема 2.3 Монтаж кабельных линий, комплектных шинопроводов и троллейных линий (часть 1)",
             "hours": 2, "is_practice": False, "is_lab": False, "practice_prep": 0,
             "equipment": BOOK_OI2, "task": "Стр. 228-241", "note": "Вх.контр"},
            {"name": "Тема 2.3 Монтаж кабельных линий, комплектных шинопроводов и троллейных линий (часть 2)",
             "hours": 2, "is_practice": False, "is_lab": False, "practice_prep": 0,
             "equipment": BOOK_OI2, "task": "Стр. 242-255", "note": "Текущий контроль"},
            {"name": "Тема 2.3 Монтаж кабельных линий, комплектных шинопроводов и троллейных линий (часть 3)",
             "hours": 2, "is_practice": False, "is_lab": False, "practice_prep": 0,
             "equipment": BOOK_OI2, "task": "Стр. 256-269", "note": "Текущий контроль"},
            {"name": "Тема 2.3 Монтаж кабельных линий, комплектных шинопроводов и троллейных линий (часть 4)",
             "hours": 2, "is_practice": False, "is_lab": False, "practice_prep": 0,
             "equipment": BOOK_OI2, "task": "Стр. 270-283", "note": "Текущий контроль"},
            # Практика (6ч = 3 занятия)
            {"name": "Практическое занятие №11 Изучение технологии выполнения разделки силового кабеля",
             "hours": 2, "is_practice": True, "is_lab": False, "practice_prep": 2,
             "equipment": PRACTICE_EQUIP, "task": "Метод. указания", "note": "Защита пр.работы"},
            {"name": "Практическое занятие №12 Изучение методов определения мест повреждения в кабельных линиях",
             "hours": 2, "is_practice": True, "is_lab": False, "practice_prep": 2,
             "equipment": PRACTICE_EQUIP, "task": "Метод. указания", "note": "Защита пр.работы"},
            {"name": "Практическое занятие №13 Расчет сечения провода по допустимой длительной токовой нагрузке и потере напряжения",
             "hours": 2, "is_practice": True, "is_lab": False, "practice_prep": 2,
             "equipment": PRACTICE_EQUIP, "task": "Метод. указания", "note": "Защита пр.работы"},
        ]
    },
    # ========= Тема 2.4 (4ч: 2 теория + 2 практика) — СЕМ.4 =========
    {
        "name": "Тема 2.4 Монтаж защитного заземления и зануления",
        "total_hours": 4,
        "practice_hours": 2,
        "semester": 4,
        "lessons": [
            {"name": "Тема 2.4 Монтаж защитного заземления и зануления",
             "hours": 2, "is_practice": False, "is_lab": False, "practice_prep": 0,
             "equipment": BOOK_OI2, "task": "Стр. 124-167", "note": "Вх.контр"},
            {"name": "Практическое занятие №14 Изучение защитного заземления, измерение сопротивления изоляции",
             "hours": 2, "is_practice": True, "is_lab": False, "practice_prep": 2,
             "equipment": PRACTICE_EQUIP, "task": "Метод. указания", "note": "Защита пр.работы"},
        ]
    },
    # ========= Тема 2.5 (16ч: 8 теория + 8 практика) — СЕМ.4 =========
    {
        "name": "Тема 2.5 Монтаж электрических машин и силовых трансформаторов",
        "total_hours": 16,
        "practice_hours": 8,
        "semester": 4,
        "lessons": [
            # Теория (8ч = 4 занятия)
            {"name": "Тема 2.5 Монтаж электрических машин и силовых трансформаторов (часть 1)",
             "hours": 2, "is_practice": False, "is_lab": False, "practice_prep": 0,
             "equipment": BOOK_OI3, "task": "Стр. 203-218", "note": "Вх.контр"},
            {"name": "Тема 2.5 Монтаж электрических машин и силовых трансформаторов (часть 2)",
             "hours": 2, "is_practice": False, "is_lab": False, "practice_prep": 0,
             "equipment": BOOK_OI3, "task": "Стр. 219-234", "note": "Текущий контроль"},
            {"name": "Тема 2.5 Монтаж электрических машин и силовых трансформаторов (часть 3)",
             "hours": 2, "is_practice": False, "is_lab": False, "practice_prep": 0,
             "equipment": BOOK_OI3, "task": "Стр. 263-273", "note": "Текущий контроль"},
            {"name": "Тема 2.5 Монтаж электрических машин и силовых трансформаторов (часть 4)",
             "hours": 2, "is_practice": False, "is_lab": False, "practice_prep": 0,
             "equipment": BOOK_OI3, "task": "Стр. 274-284", "note": "Текущий контроль"},
            # Практика (8ч = 4 занятия)
            {"name": "Практическое занятие №15 Схемы подключения (часть 1)",
             "hours": 2, "is_practice": True, "is_lab": False, "practice_prep": 2,
             "equipment": PRACTICE_EQUIP, "task": "Метод. указания", "note": "Защита пр.работы"},
            {"name": "Практическое занятие №15 Схемы подключения (часть 2)",
             "hours": 2, "is_practice": True, "is_lab": False, "practice_prep": 2,
             "equipment": PRACTICE_EQUIP, "task": "Метод. указания", "note": "Защита пр.работы"},
            {"name": "Практическое занятие №16 Расчет основных параметров трехфазного трансформатора (часть 1)",
             "hours": 2, "is_practice": True, "is_lab": False, "practice_prep": 2,
             "equipment": PRACTICE_EQUIP, "task": "Метод. указания", "note": "Защита пр.работы"},
            {"name": "Практическое занятие №16 Расчет основных параметров трехфазного трансформатора (часть 2)",
             "hours": 2, "is_practice": True, "is_lab": False, "practice_prep": 2,
             "equipment": PRACTICE_EQUIP, "task": "Метод. указания", "note": "Защита пр.работы"},
        ]
    },
    # ========= Тема 2.6 (10ч: 6 теория + 4 практика) — СЕМ.4 =========
    {
        "name": "Тема 2.6 Ремонт осветительных электроустановок, аппаратов защиты, пускорегулирующей аппаратуры",
        "total_hours": 10,
        "practice_hours": 4,
        "semester": 4,
        "lessons": [
            # Теория (6ч = 3 занятия)
            {"name": "Тема 2.6 Ремонт осветительных электроустановок, аппаратов защиты, пускорегулирующей аппаратуры (часть 1)",
             "hours": 2, "is_practice": False, "is_lab": False, "practice_prep": 0,
             "equipment": BOOK_OI3, "task": "Стр. 175-181", "note": "Вх.контр"},
            {"name": "Тема 2.6 Ремонт осветительных электроустановок, аппаратов защиты, пускорегулирующей аппаратуры (часть 2)",
             "hours": 2, "is_practice": False, "is_lab": False, "practice_prep": 0,
             "equipment": BOOK_OI3, "task": "Стр. 182-188", "note": "Текущий контроль"},
            {"name": "Тема 2.6 Ремонт осветительных электроустановок, аппаратов защиты, пускорегулирующей аппаратуры (часть 3)",
             "hours": 2, "is_practice": False, "is_lab": False, "practice_prep": 0,
             "equipment": BOOK_OI3, "task": "Стр. 189-194", "note": "Текущий контроль"},
            # Практика (4ч = 2 занятия)
            {"name": "Практическое занятие №17 Устранение неисправностей в электрической схеме пуска и реверса электрического двигателя с короткозамкнутым ротором (часть 1)",
             "hours": 2, "is_practice": True, "is_lab": False, "practice_prep": 2,
             "equipment": PRACTICE_EQUIP, "task": "Метод. указания", "note": "Защита пр.работы"},
            {"name": "Практическое занятие №17 Устранение неисправностей в электрической схеме пуска и реверса электрического двигателя с короткозамкнутым ротором (часть 2)",
             "hours": 2, "is_practice": True, "is_lab": False, "practice_prep": 2,
             "equipment": PRACTICE_EQUIP, "task": "Метод. указания", "note": "Защита пр.работы"},
        ]
    },
    # ========= Тема 2.7 (10ч: 6 теория + 4 практика) — СЕМ.4 =========
    {
        "name": "Тема 2.7 Ремонт воздушных и кабельных линий электропередачи",
        "total_hours": 10,
        "practice_hours": 4,
        "semester": 4,
        "lessons": [
            # Теория (6ч = 3 занятия)
            {"name": "Тема 2.7 Ремонт воздушных и кабельных линий электропередачи (часть 1)",
             "hours": 2, "is_practice": False, "is_lab": False, "practice_prep": 0,
             "equipment": BOOK_OI3, "task": "Стр. 285-296", "note": "Вх.контр"},
            {"name": "Тема 2.7 Ремонт воздушных и кабельных линий электропередачи (часть 2)",
             "hours": 2, "is_practice": False, "is_lab": False, "practice_prep": 0,
             "equipment": BOOK_OI3, "task": "Стр. 297-308", "note": "Текущий контроль"},
            {"name": "Тема 2.7 Ремонт воздушных и кабельных линий электропередачи (часть 3)",
             "hours": 2, "is_practice": False, "is_lab": False, "practice_prep": 0,
             "equipment": BOOK_OI3, "task": "Стр. 309-320", "note": "Текущий контроль"},
            # Практика (4ч = 2 занятия)
            {"name": "Практическое занятие №18 Определение основных неисправностей в кабельных и воздушных линиях электропередач и способы их устранения (часть 1)",
             "hours": 2, "is_practice": True, "is_lab": False, "practice_prep": 2,
             "equipment": PRACTICE_EQUIP, "task": "Метод. указания", "note": "Защита пр.работы"},
            {"name": "Практическое занятие №18 Определение основных неисправностей в кабельных и воздушных линиях электропередач и способы их устранения (часть 2)",
             "hours": 2, "is_practice": True, "is_lab": False, "practice_prep": 2,
             "equipment": PRACTICE_EQUIP, "task": "Метод. указания", "note": "Защита пр.работы"},
        ]
    },
    # ========= Тема 2.8 (22ч: 6 теория + 16 практика) — СЕМ.4 =========
    {
        "name": "Тема 2.8 Ремонт электрических машин и трансформаторов",
        "total_hours": 22,
        "practice_hours": 16,
        "semester": 4,
        "lessons": [
            # Теория (6ч = 3 занятия)
            {"name": "Тема 2.8 Ремонт электрических машин и трансформаторов (часть 1)",
             "hours": 2, "is_practice": False, "is_lab": False, "practice_prep": 0,
             "equipment": BOOK_OI3, "task": "Стр. 203-218", "note": "Вх.контр"},
            {"name": "Тема 2.8 Ремонт электрических машин и трансформаторов (часть 2)",
             "hours": 2, "is_practice": False, "is_lab": False, "practice_prep": 0,
             "equipment": BOOK_OI3, "task": "Стр. 219-234", "note": "Текущий контроль"},
            {"name": "Тема 2.8 Ремонт электрических машин и трансформаторов (часть 3)",
             "hours": 2, "is_practice": False, "is_lab": False, "practice_prep": 0,
             "equipment": BOOK_OI3, "task": "Стр. 263-284", "note": "Текущий контроль"},
            # Практика (16ч = 8 занятий)
            {"name": "Практическое занятие №19 Исследование асинхронного двигателя с короткозамкнутым ротором (часть 1)",
             "hours": 2, "is_practice": True, "is_lab": False, "practice_prep": 2,
             "equipment": PRACTICE_EQUIP, "task": "Метод. указания", "note": "Защита пр.работы"},
            {"name": "Практическое занятие №19 Исследование асинхронного двигателя с короткозамкнутым ротором (часть 2)",
             "hours": 2, "is_practice": True, "is_lab": False, "practice_prep": 2,
             "equipment": PRACTICE_EQUIP, "task": "Метод. указания", "note": "Защита пр.работы"},
            {"name": "Практическое занятие №20 Прямой пуск в ход трехфазного асинхронного двигателя (часть 1)",
             "hours": 2, "is_practice": True, "is_lab": False, "practice_prep": 2,
             "equipment": PRACTICE_EQUIP, "task": "Метод. указания", "note": "Защита пр.работы"},
            {"name": "Практическое занятие №20 Прямой пуск в ход трехфазного асинхронного двигателя (часть 2)",
             "hours": 2, "is_practice": True, "is_lab": False, "practice_prep": 2,
             "equipment": PRACTICE_EQUIP, "task": "Метод. указания", "note": "Защита пр.работы"},
            {"name": "Практическое занятие №21 Определение внешней характеристики, группы соединения обмоток и коэффициента трансформации трансформатора (часть 1)",
             "hours": 2, "is_practice": True, "is_lab": False, "practice_prep": 2,
             "equipment": PRACTICE_EQUIP, "task": "Метод. указания", "note": "Защита пр.работы"},
            {"name": "Практическое занятие №21 Определение внешней характеристики, группы соединения обмоток и коэффициента трансформации трансформатора (часть 2)",
             "hours": 2, "is_practice": True, "is_lab": False, "practice_prep": 2,
             "equipment": PRACTICE_EQUIP, "task": "Метод. указания", "note": "Защита пр.работы"},
            {"name": "Практическое занятие №22 Определение основных неисправностей электрических машин и трансформаторов и способы их устранения (часть 1)",
             "hours": 2, "is_practice": True, "is_lab": False, "practice_prep": 2,
             "equipment": PRACTICE_EQUIP, "task": "Метод. указания", "note": "Защита пр.работы"},
            {"name": "Практическое занятие №22 Определение основных неисправностей электрических машин и трансформаторов и способы их устранения (часть 2)",
             "hours": 2, "is_practice": True, "is_lab": False, "practice_prep": 2,
             "equipment": PRACTICE_EQUIP, "task": "Метод. указания", "note": "Защита пр.работы"},
        ]
    },
    # ========= Дифференцированный зачет (2ч) — СЕМ.4 =========
    {
        "name": "Дифференцированный зачет",
        "total_hours": 2,
        "practice_hours": 0,
        "semester": 4,
        "lessons": [
            {"name": "Дифференцированный зачет",
             "hours": 2, "is_practice": False, "is_lab": False,
             "practice_prep": 0,
             "equipment": "",
             "task": "",
             "note": "Промежуточная аттестация",
             "is_exam": True},
        ]
    },
]

# ============================================================
# МАТЕРИАЛЬНО-ТЕХНИЧЕСКОЕ ОБЕСПЕЧЕНИЕ
# ============================================================
EQUIPMENT = [
    "Кабинет физики, электротехники и электроники",
    "Лаборатория электротехники и электроники",
    "Методические указания по выполнению лабораторно-практических работ",
    "Лаборатория электрических измерений и электрических цепей",
    "Мастерская \"Слесарно-механическая\"",
    "Мастерская \"Электротехническая\"",
    "Мастерская \"Монтажа, технического обслуживания и эксплуатации электрооборудования\"",
    "Средства мультимедиа",
    "ЭВМ",
]

# ============================================================
# ИНФОРМАЦИОННОЕ ОБЕСПЕЧЕНИЕ
# ============================================================
SOURCES_BASIC = [
    {"code": "ОИ 1", "name": "Общая технология электромонтажных работ: учебник для СПО",
     "author": "Григорьева С.В.", "publisher": "М.: ИЦ \"Академия\", 2020"},
    {"code": "ОИ 2", "name": "Технология электромонтажных работ: учеб. пособие для СПО",
     "author": "Нестеренко В.М.", "publisher": "М.: ИЦ \"Академия\", 2022"},
    {"code": "ОИ 3", "name": "Сборка, монтаж, регулировка и ремонт узлов и механизмов оборудования, агрегатов, машин, станков и другого электрооборудования промышленных организаций: учебник",
     "author": "Сидорова Л.Г.", "publisher": "М.: ИЦ \"Академия\", 2022"},
    {"code": "ОИ 4", "name": "Проверка и наладка электрооборудования: учебник",
     "author": "Ярочкина Г.В.", "publisher": "М.: ИЦ \"Академия\", 2022"},
    {"code": "ОИ 5", "name": "Организация и выполнение работ по монтажу и наладке электрооборудования промышленных и гражданских зданий. В двух частях. Часть 1. Внутреннее электроснабжение промышленных и гражданских зданий: учебник",
     "author": "Бычков А.В.", "publisher": "М.: ИЦ \"Академия\", 2020"},
    {"code": "ОИ 6", "name": "Организация и выполнение работ по монтажу и наладке электрооборудования промышленных и гражданских зданий. В двух частях. Часть 2. Монтаж и наладка электрооборудования промышленных и гражданских зданий: учебник",
     "author": "Шашкова И.В., Бычков А.В.", "publisher": "М.: ИЦ \"Академия\", 2020"},
    {"code": "ОИ 7", "name": "Монтаж, наладка, эксплуатация и ремонт систем электроснабжения промышленных предприятий: учебное пособие для СПО",
     "author": "Полуянович Н.К.", "publisher": "Санкт-Петербург: Лань, 2022"},
]

SOURCES_ADDITIONAL = [
    {"code": "ДИ 1", "name": "Электрические системы и сети. Энергосбережение: учебное пособие для СПО",
     "author": "Климова Г.Н.", "publisher": "М.: Издательство Юрайт, 2023"},
    {"code": "ДИ 2", "name": "Организация и методика производственного обучения. Электромонтер-кабельщик: учебное пособие для СПО",
     "author": "Бредихин А.Н.", "publisher": "М.: Издательство Юрайт, 2023"},
    {"code": "ДИ 3", "name": "Информационный портал для электромонтеров",
     "author": "", "publisher": "http://electromonter.info"},
    {"code": "ДИ 4", "name": "Образовательный сайт \"Школа для электрика\"",
     "author": "", "publisher": "http://ElectricalSchool.info"},
    {"code": "ДИ 5", "name": "Нормативно-технические документы",
     "author": "", "publisher": "http://electrolibrary.info"},
]


# ============================================================
# ГЕНЕРАЦИЯ DOCX
# ============================================================

def generate_ktp(output_path):
    doc = Document()
    data = MDK_DATA

    # МАРГИНЫ
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(1.5)

    # ==================== ТИТУЛЬНЫЙ ЛИСТ ====================
    add_paragraph(doc, 'МИНИСТЕРСТВО ОБЩЕГО И ПРОФЕССИОНАЛЬНОГО ОБРАЗОВАНИЯ РОСТОВСКОЙ ОБЛАСТИ', bold=True, size=11)
    add_paragraph(doc, '')
    add_paragraph(doc, 'ГОСУДАРСТВЕННОЕ БЮДЖЕТНОЕ ПРОФЕССИОНАЛЬНОЕ ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ РОСТОВСКОЙ ОБЛАСТИ', bold=True, size=11)
    add_paragraph(doc, '\u00abСАЛЬСКИЙ ИНДУСТРИАЛЬНЫЙ ТЕХНИКУМ\u00bb', bold=True, size=12)
    add_paragraph(doc, '(ГБПОУ РО \u00abСИТ\u00bb)', size=11)
    add_paragraph(doc, '')
    add_paragraph(doc, '')

    # Таблица утверждения
    t = doc.add_table(rows=3, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.RIGHT
    set_cell_text(t.cell(0, 1), 'УТВЕРЖДАЮ', bold=True, size=10)
    set_cell_text(t.cell(1, 1), 'Зам. директора ГБПОУ РО \u00abСИТ\u00bb', size=10)
    set_cell_text(t.cell(1, 2), f'____________ {data["director"]}', size=10)
    set_cell_text(t.cell(2, 1), '\u00ab____\u00bb ______________ 20 ____ г.', size=10)
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)

    add_paragraph(doc, '')
    add_paragraph(doc, 'КАЛЕНДАРНО-ТЕМАТИЧЕСКИЙ ПЛАН', bold=True, size=14)
    add_paragraph(doc, '')

    add_paragraph(doc, f'на 3, 4 семестры {data["year"]} учебного года {data["course"]} курс',
                  size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, '')
    add_paragraph(doc, f'учебной группы (учебных групп) {data["group"]}',
                  size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, '')
    add_paragraph(doc, f'Профессиональный модуль: ПМ.{data["pm_code"]} {data["pm_name"]}',
                  size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, '')
    add_paragraph(doc, f'Междисциплинарные курсы: МДК {data["mdk_code"]} {data["mdk_name"]}',
                  size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, '')
    add_paragraph(doc, f'по профессии: {data["profession"]}',
                  size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, '')

    total = data['total_hours']
    practice_total = data['practice_hours']
    with_teacher = data.get('with_teacher_hours', total)
    theory = data.get('theory_hours', 0)
    practical = data.get('practical_hours', 0)
    lab = data.get('lab_hours', 0)
    coursework = data.get('coursework_hours', 0)
    independent = data.get('independent_hours', 0)
    exam = data.get('exam_form', '')

    add_paragraph(doc, f'Объем образовательной программы: ______{total}______________ (часов);',
                  size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, f'в том числе в форме практической подготовки ___{practice_total}______________(часа);',
                  size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, '')
    add_paragraph(doc, f'Учебная нагрузка во взаимодействии с преподавателем ________{with_teacher}_______________ (часов):',
                  size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, 'из нее:', size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, f'теоретическое обучение __{theory}___ (часов);                    практические занятия __{practical}__ (часов);',
                  size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, f'лабораторные занятия ___{"____" if not lab else lab}____ (часов);                   курсовая работа/проект ___{"____" if not coursework else coursework}____ (часов);',
                  size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, f'самостоятельная работа ___{"____" if not independent else independent}____ (часов);',
                  size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, f'промежуточная аттестация в форме ___{exam}___',
                  size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, '                                                                               (указать форму)',
                  size=9, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, '')

    chair = data.get('commission_chair', 'Ткаченко А.Н.')
    add_paragraph(doc, f'Составлен в соответствии с рабочей программой ПМ {data["pm_code"]}, утверждённой __________________',
                  size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, '              (дата утверждения)', size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, '')
    add_paragraph(doc, 'Рассмотрен на заседании цикловой комиссии _____________________________ дисциплин',
                  size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, 'Протокол  №_______от________________20_____ года',
                  size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, f'Председатель цикловой комиссии ___________________________/{chair}/',
                  size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, '')
    add_paragraph(doc, '')
    add_paragraph(doc, 'г. Сальск', size=11, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph(doc, f'{data["year"]} уч. год', size=11, alignment=WD_ALIGN_PARAGRAPH.RIGHT)

    # ==================== ТАБЛИЦА 1 ====================
    doc.add_page_break()
    add_paragraph(doc, 'Распределение часов по профессиональному модулю', bold=True, size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, 'Таблица 1', bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, '')

    t1_headers = ['Междисциплинарный курс (индекс МДК)', 'Курс', 'Семестр',
                   'Объём образовательной программы', 'Всего, часов',
                   'Теоретические занятия', 'Лабораторные работы, часов',
                   'Практические занятия, часов', 'Курсовые работы (проекты), часов',
                   'Самостоятельная работа обучающегося, часов',
                   'Учебная практика, часов', 'Производственная практика, часов']

    semesters = data.get('semesters', [])
    num_sem_rows = len(semesters) if semesters else 1
    t1 = doc.add_table(rows=3 + num_sem_rows, cols=len(t1_headers))
    t1.style = 'Table Grid'

    for j, h in enumerate(t1_headers):
        set_cell_text(t1.cell(0, j), h, bold=True, size=7, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(t1.cell(0, j), 'D9E2F3')

    for si, sem in enumerate(semesters):
        row = t1.rows[1 + si]
        set_cell_text(row.cells[0], f'МДК {data["mdk_code"]}', size=8)
        set_cell_text(row.cells[1], str(data['course']), size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[2], str(sem['sem']), size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[3], str(sem['total']), size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[4], str(sem['total']), size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[5], str(sem['theory']), size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[6], str(lab) if lab else '-', size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[7], str(sem['practice']), size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[8], str(coursework) if coursework else '-', size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[9], str(independent) if independent else '-', size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[10], '-', size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[11], '-', size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    row_p = t1.rows[1 + num_sem_rows]
    set_cell_text(row_p.cells[0], 'Практика', size=8)
    for j in range(len(t1_headers)):
        if j not in [0]:
            set_cell_text(row_p.cells[j], '-', size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    row_t = t1.rows[2 + num_sem_rows]
    set_cell_text(row_t.cells[0], 'Всего', size=8, bold=True)
    set_cell_text(row_t.cells[3], str(total), size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    set_cell_text(row_t.cells[4], str(with_teacher), size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    set_cell_text(row_t.cells[5], str(theory), size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    set_cell_text(row_t.cells[6], str(lab) if lab else '-', size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    set_cell_text(row_t.cells[7], str(practical) if practical else '-', size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    set_cell_text(row_t.cells[8], str(coursework) if coursework else '-', size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    set_cell_text(row_t.cells[9], str(independent) if independent else '-', size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    set_cell_text(row_t.cells[10], '-', size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    set_cell_text(row_t.cells[11], '-', size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

    add_paragraph(doc, '')
    add_paragraph(doc, 'Форма промежуточной аттестации обучающихся за семестр по междисциплинарному курсу',
                  size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, f'(МДК {data["mdk_code"]} {data["mdk_name"]}) \u2013 ({exam}).',
                  size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    # ==================== ТАБЛИЦА 2 (11 колонок) ====================
    add_paragraph(doc, '')
    add_paragraph(doc, 'Содержание обучения по профессиональному модулю', bold=True, size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, '                                                                                                                                        Таблица 2',
                  size=10, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
    add_paragraph(doc, '')

    # Ширины колонок (EMU) — по эталону КТП МДК 02.01
    T2_COL_WIDTHS = [
        318135,   # 0: № занятия
        1862455,  # 1: Наименование
        1403985,  # 2: Кол-во часов
        1403985,  # 3: В форме практической подготовки
        1403985,  # 4: Вид занятия
        937895,   # 5: ОК
        937895,   # 6: ПК
        577215,   # 7: Материальное и информационное обеспечение
        585470,   # 8: Задания для студентов
        600710,   # 9: Формы и методы контроля
        543560,   # 10: ФИО преподавателя
    ]

    # Считаем строки: заголовок(4) + МДК(1) + темы(N) + занятия + Итого(1)
    num_lessons = sum(len(topic['lessons']) for topic in TOPICS)
    num_topic_rows = len(TOPICS)
    total_rows = 4 + 1 + num_topic_rows + num_lessons + 1

    t2 = doc.add_table(rows=total_rows, cols=11)
    t2.style = 'Table Grid'

    # --- Заголовок Таблицы 2 (3 строки + номерная строка) ---

    # Row 0: основные заголовки (с мержами)
    # Столбцы: 0=№, 1=Наименование, 2-4=Обязательная нагрузка (gridSpan 3),
    #          5-6=Компетенции (gridSpan 2), 7=Обеспечение, 8=Задания, 9=Контроль, 10=ФИО
    r = t2.rows[0]
    set_cell_text(r.cells[0], '№ занятия', bold=True, size=7, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(r.cells[0], 'D9E2F3')
    set_cell_text(r.cells[1], 'Наименование разделов\nпрофессионального модуля,\nтем и занятий по МДК', bold=True, size=7, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(r.cells[1], 'D9E2F3')
    # Обязательная нагрузка — мерж 3 ячеек
    r.cells[2].merge(r.cells[4])
    set_cell_text(r.cells[2], 'Обязательная учебная нагрузка', bold=True, size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(r.cells[2], 'D9E2F3')
    # Компетенции — мерж 2 ячеек
    r.cells[5].merge(r.cells[6])
    set_cell_text(r.cells[5], 'Коды формируемых компетенции', bold=True, size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(r.cells[5], 'D9E2F3')
    set_cell_text(r.cells[7], 'Материальное и информационное\nобеспечение занятий', bold=True, size=7, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(r.cells[7], 'D9E2F3')
    set_cell_text(r.cells[8], 'Задания для студентов', bold=True, size=7, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(r.cells[8], 'D9E2F3')
    set_cell_text(r.cells[9], 'Формы и методы контроля', bold=True, size=7, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(r.cells[9], 'D9E2F3')
    set_cell_text(r.cells[10], 'ФИО преподавателя', bold=True, size=7, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(r.cells[10], 'D9E2F3')

    # Row 1: подзаголовки для «Обязательная нагрузка» и «Компетенции»
    r1 = t2.rows[1]
    # Ячейки 0,1,7,8,9,10 — пустые (vMerge продолжение)
    for ci in [0, 1, 7, 8, 9, 10]:
        set_cell_shading(r1.cells[ci], 'D9E2F3')
    # Кол-во часов
    set_cell_text(r1.cells[2], 'Кол-во\nчасов', bold=True, size=7, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(r1.cells[2], 'D9E2F3')
    # В форме практической подготовки
    set_cell_text(r1.cells[3], 'В форме практической\nподготовки', bold=True, size=7, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(r1.cells[3], 'D9E2F3')
    # Вид занятия
    set_cell_text(r1.cells[4], 'Вид занятия', bold=True, size=7, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(r1.cells[4], 'D9E2F3')
    # Компетенции — мерж 2 ячеек (продолжение)
    r1.cells[5].merge(r1.cells[6])
    set_cell_shading(r1.cells[5], 'D9E2F3')

    # Row 2: ОК, ПК
    r2 = t2.rows[2]
    for ci in [0, 1, 2, 3, 4, 7, 8, 9, 10]:
        set_cell_shading(r2.cells[ci], 'D9E2F3')
    set_cell_text(r2.cells[5], 'ОК', bold=True, size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(r2.cells[5], 'D9E2F3')
    set_cell_text(r2.cells[6], 'ПК', bold=True, size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_shading(r2.cells[6], 'D9E2F3')

    # Row 3: номера граф
    r3 = t2.rows[3]
    for ci, num in enumerate(['1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11']):
        set_cell_text(r3.cells[ci], num, bold=True, size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(r3.cells[ci], 'D9E2F3')

    # Установка ширин всех ячеек заголовка
    for ri in range(4):
        for ci in range(11):
            set_col_width(t2.rows[ri].cells[ci], T2_COL_WIDTHS[ci])

    # --- Данные ---
    row_idx = 4

    # Строка МДК
    set_cell_text(t2.cell(row_idx, 1), f'МДК {data["mdk_code"]}   {data["mdk_name"]}', bold=True, size=8)
    set_cell_text(t2.cell(row_idx, 2), str(total), bold=True, size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(t2.cell(row_idx, 3), str(practice_total), bold=True, size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for j in range(11):
        set_cell_shading(t2.cell(row_idx, j), 'F2F2F2')
        set_col_width(t2.cell(row_idx, j), T2_COL_WIDTHS[j])
    row_idx += 1

    lesson_num = 1
    for topic in TOPICS:
        # Заголовочная строка темы
        set_cell_text(t2.cell(row_idx, 1), topic['name'], bold=True, size=8)
        set_cell_text(t2.cell(row_idx, 2), str(topic['total_hours']), bold=True, size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(t2.cell(row_idx, 3), str(topic['practice_hours']) if topic['practice_hours'] else '',
                      bold=True, size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        for j in range(11):
            set_cell_shading(t2.cell(row_idx, j), 'F2F2F2')
            set_col_width(t2.cell(row_idx, j), T2_COL_WIDTHS[j])
        row_idx += 1

        # Занятия
        for lesson in topic['lessons']:
            # col 0: № занятия
            set_cell_text(t2.cell(row_idx, 0), str(lesson_num), size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            # col 1: Наименование
            set_cell_text(t2.cell(row_idx, 1), lesson['name'], size=8)
            # col 2: Кол-во часов
            set_cell_text(t2.cell(row_idx, 2), str(lesson['hours']), size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            # col 3: В форме практической подготовки
            if lesson.get('practice_prep'):
                set_cell_text(t2.cell(row_idx, 3), str(lesson['practice_prep']), size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            # col 4: Вид занятия
            if lesson.get('is_exam'):
                vid = 'Диф. зачет'
            elif lesson.get('is_lab'):
                vid = 'Лаб. работа'
            elif lesson.get('is_practice'):
                vid = 'Практ. занятие'
            else:
                vid = 'Урок'
            set_cell_text(t2.cell(row_idx, 4), vid, size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            # col 5: ОК (пусто)
            # col 6: ПК (пусто)
            # col 7: Материальное и информационное обеспечение
            set_cell_text(t2.cell(row_idx, 7), lesson.get('equipment', ''), size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            # col 8: Задания для студентов
            set_cell_text(t2.cell(row_idx, 8), lesson.get('task', ''), size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            # col 9: Формы и методы контроля
            set_cell_text(t2.cell(row_idx, 9), lesson.get('note', ''), size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            # col 10: ФИО преподавателя (пусто)

            # Установка ширин
            for j in range(11):
                set_col_width(t2.cell(row_idx, j), T2_COL_WIDTHS[j])

            lesson_num += 1
            row_idx += 1

    # Итого
    set_cell_text(t2.cell(row_idx, 1), 'Итого:', bold=True, size=8)
    set_cell_text(t2.cell(row_idx, 2), str(total), bold=True, size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(t2.cell(row_idx, 3), str(practice_total), bold=True, size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for j in range(11):
        set_cell_shading(t2.cell(row_idx, j), 'D9E2F3')
        set_col_width(t2.cell(row_idx, j), T2_COL_WIDTHS[j])

    # ==================== ТАБЛИЦА 2а ====================
    doc.add_page_break()
    add_paragraph(doc, 'Материально-техническое обеспечение занятий', bold=True, size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, 'Таблица 2а', bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT)
    add_paragraph(doc, '')

    t3 = doc.add_table(rows=max(len(EQUIPMENT), 1) + 1, cols=2)
    t3.style = 'Table Grid'
    set_cell_text(t3.cell(0, 0), '№ п/п', bold=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(t3.cell(0, 1), 'Материально-техническое обеспечение занятий', bold=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for j in range(2):
        set_cell_shading(t3.cell(0, j), 'D9E2F3')
    for i, eq in enumerate(EQUIPMENT):
        set_cell_text(t3.cell(i + 1, 0), str(i + 1), size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(t3.cell(i + 1, 1), eq, size=9)

    # ==================== ТАБЛИЦЫ 2б и 2в ====================
    add_paragraph(doc, '')
    add_paragraph(doc, 'Информационное обеспечение обучения', bold=True, size=11, alignment=WD_ALIGN_PARAGRAPH.LEFT)

    if SOURCES_BASIC:
        add_paragraph(doc, 'Основные источники (ОИ):', bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        add_paragraph(doc, 'Таблица 2б', bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        add_paragraph(doc, '')
        t4 = doc.add_table(rows=len(SOURCES_BASIC) + 1, cols=4)
        t4.style = 'Table Grid'
        for j, h in enumerate(['№ п/п', 'Наименование', 'Автор', 'Издательство, год издания']):
            set_cell_text(t4.cell(0, j), h, bold=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_shading(t4.cell(0, j), 'D9E2F3')
        for i, src in enumerate(SOURCES_BASIC):
            set_cell_text(t4.cell(i + 1, 0), src.get('code', f'ОИ {i+1}'), size=9)
            set_cell_text(t4.cell(i + 1, 1), src.get('name', ''), size=9)
            set_cell_text(t4.cell(i + 1, 2), src.get('author', ''), size=9)
            set_cell_text(t4.cell(i + 1, 3), src.get('publisher', ''), size=9)

    if SOURCES_ADDITIONAL:
        add_paragraph(doc, '')
        add_paragraph(doc, 'Дополнительные источники (ДИ):', bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        add_paragraph(doc, 'Таблица 2в', bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.LEFT)
        add_paragraph(doc, '')
        t5 = doc.add_table(rows=len(SOURCES_ADDITIONAL) + 1, cols=4)
        t5.style = 'Table Grid'
        for j, h in enumerate(['№ п/п', 'Наименование', 'Автор', 'Издательство, год издания']):
            set_cell_text(t5.cell(0, j), h, bold=True, size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_shading(t5.cell(0, j), 'D9E2F3')
        for i, src in enumerate(SOURCES_ADDITIONAL):
            set_cell_text(t5.cell(i + 1, 0), src.get('code', f'ДИ {i+1}'), size=9)
            set_cell_text(t5.cell(i + 1, 1), src.get('name', ''), size=9)
            set_cell_text(t5.cell(i + 1, 2), src.get('author', ''), size=9)
            set_cell_text(t5.cell(i + 1, 3), src.get('publisher', ''), size=9)

    doc.save(output_path)
    print(f'КТП сохранён: {output_path}')

    # ==================== ПРОВЕРКА ====================
    total_lessons = sum(len(t['lessons']) for t in TOPICS)
    total_hours_check = sum(l['hours'] for t in TOPICS for l in t['lessons'])
    total_practice_check = sum(l.get('practice_prep', 0) for t in TOPICS for l in t['lessons'])
    theory_check = sum(l['hours'] for t in TOPICS for l in t['lessons'] if not l.get('is_practice') and not l.get('is_lab'))
    practice_check = sum(l['hours'] for t in TOPICS for l in t['lessons'] if l.get('is_practice') or l.get('is_lab'))

    print(f'\n=== ПРОВЕРКА ===')
    print(f'Всего занятий: {total_lessons}')
    print(f'Сумма часов: {total_hours_check} (ожидается {total}) -> {"OK" if total_hours_check == total else "ОШИБКА"} ')
    print(f'Практ. подготовка: {total_practice_check} (ожидается {practice_total}) -> {"OK" if total_practice_check == practice_total else "ОШИБКА"} ')
    print(f'Теория: {theory_check} (ожидается {theory}) -> {"OK" if theory_check == theory else "ОШИБКА"} ')
    print(f'Практика: {practice_check} (ожидается {practical}) -> {"OK" if practice_check == practical else "ОШИБКА"} ')

    # Проверка по семестрам
    sem3_h = 0
    sem4_h = 0
    running = 0
    for t in TOPICS:
        sem = t.get('semester', 4)
        for l in t['lessons']:
            running += l['hours']
            if isinstance(sem, list):
                if running <= 36:
                    sem3_h += l['hours']
                else:
                    sem4_h += l['hours']
            elif sem == 3:
                sem3_h += l['hours']
            else:
                sem4_h += l['hours']
    print(f'Сем.3: {sem3_h}ч (ожидается 36) -> {"OK" if sem3_h == 36 else "ОШИБКА"} ')
    print(f'Сем.4: {sem4_h}ч (ожидается 76) -> {"OK" if sem4_h == 76 else "ОШИБКА"} ')

    # Проверка столбцов 8 и 9
    print(f'\n=== ПРОВЕРКА СТОЛБЦОВ 8 И 9 ===')
    for topic in TOPICS:
        for l in topic['lessons']:
            eq = l.get('equipment', '')
            task = l.get('task', '')
            is_p = l.get('is_practice', False)
            if is_p:
                ok_eq = (eq == PRACTICE_EQUIP)
                ok_task = (task == 'Метод. указания')
            else:
                ok_eq = eq in (BOOK_OI2, BOOK_OI3)
                ok_task = task.startswith('Стр. ')
            status = 'OK' if (ok_eq and ok_task) else 'ОШИБКА'
            if status == 'ОШИБКА':
                print(f'  {status}: {l["name"][:50]}... eq={eq!r} task={task!r}')
    print('Все строки столбцов 8-9 проверены.')


if __name__ == '__main__':
    output = '/home/z/my-project/YMK/KTP/МДК_05_02_КТП.docx'
    generate_ktp(output)
