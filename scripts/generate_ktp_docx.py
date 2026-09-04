#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Генератор КТП в формате DOCX — СТРОГО по единственному эталону проекта:
KTP/КТП МДК 05.02 М2 курс.docx

Использование:
    python3 scripts/generate_ktp_docx.py --data data.json --output "KTP/КТП МДК 05.02 М2 курс.docx"

Структура документа (7 таблиц, как в эталоне):
  1. Титульный лист (министерство, УТВЕРЖДАЮ, шапка КТП, часы, цикловая комиссия)
  2. Таблица 1  — Распределение часов по профессиональному модулю
                   (12 колонок, 4-строчная шапка с объединениями, 12 строк)
  3. Таблица 2  — Содержание обучения
                   (11 колонок, шапка 3 строки + номерная, объединения gridSpan/vMerge)
  4. Таблица 2а — Материально-техническое обеспечение (2 колонки + номерная строка)
  5. Таблица 2б — Основные источники / ОИ (4 колонки)
  6. Таблица 2в — Дополнительные источники / ДИ (4 колонки)
  7. Таблица 2г — Электронные источники / ЭИ (5 колонок, включая URL)

Требования эталона, зашитые в код:
  - «МИНИСТЕРСТВО ОБРАЗОВАНИЯ РОСТОВСКОЙ ОБЛАСТИ» (без «общего и профессионального»)
  - БЕЗ заливки ячеек (все ячейки белые)
  - Шрифт Times New Roman: титульник 12pt, таблицы 10pt (наименования литературы 9pt)
  - Выравнивание в ячейках по центру (кроме наименований — по левому краю)
  - Вертикальное выравнивание ячеек — по центру
  - Строка промежуточной аттестации в Таблице 2: номер, название, 2 часа,
    остальные графы ПУСТЫЕ (строго по эталону)
  - Заливки D9E2F3/F2F2F2 НЕ используются (в эталоне их нет)

Формат входных данных: см. scripts/data_05_02_example.json и README.md
"""

import argparse
import json
import sys

from docx import Document
from docx.shared import Pt, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = 'Times New Roman'


def hour_word(n):
    """Склонение «час» после числа: 1→час, 2–4→часа, 5–20 и десятки→часов.
    Примеры: 8→часов, 2→часа, 60→часов, 112→часов, 22→часа, 12→часов."""
    if n in (None, '', 0):
        return 'часов'
    n = int(n)
    if n % 100 in (11, 12, 13, 14):
        return 'часов'
    if n % 10 == 1:
        return 'час'
    if n % 10 in (2, 3, 4):
        return 'часа'
    return 'часов'

# ======================================================================
# Утилиты низкого уровня
# ======================================================================

def _fmt_run(run, size=12, bold=False):
    """Единая настройка шрифта run'а (ascii/hAnsi/eastAsia/cs)."""
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), FONT)
    rFonts.set(qn('w:hAnsi'), FONT)
    rFonts.set(qn('w:eastAsia'), FONT)
    rFonts.set(qn('w:cs'), FONT)


def cell_text(cell, text, bold=False, size=10, align=WD_ALIGN_PARAGRAPH.CENTER):
    """Заполнить ячейку текстом (многострочным через \\n), шрифт по эталону."""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    lines = str(text).split('\n')
    for i, line in enumerate(lines):
        run = p.add_run(line)
        _fmt_run(run, size=size, bold=bold)
        if i < len(lines) - 1:
            run.add_break()
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_p(doc, text='', style=None, align=WD_ALIGN_PARAGRAPH.LEFT,
          size=12, bold=False, space_after=0):
    """Обычный абзац."""
    p = doc.add_paragraph(style=style)
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.0
    if text:
        run = p.add_run(text)
        _fmt_run(run, size=size, bold=bold)
    return p


def add_fillin_p(doc, segments, style=None, align=WD_ALIGN_PARAGRAPH.LEFT,
                 size=12, bold=False, space_after=0):
    """Абзац с заполняемыми полями: segments = [(текст, underline), ...]."""
    p = doc.add_paragraph(style=style)
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.0
    for text, underline in segments:
        run = p.add_run(text)
        _fmt_run(run, size=size, bold=bold)
        run.font.underline = underline
    return p


def page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    return p


def new_table(doc, rows, cols, widths, bordered=True):
    """Таблица с фиксированной раскладкой и ширинами колонок (twips)."""
    t = doc.add_table(rows=rows, cols=cols)
    if bordered:
        t.style = 'Table Grid'
    # без стиля таблица рендерится в Word без границ (нужно для УТВЕРЖДАЮ)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    tblPr = t._tbl.tblPr
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)
    for i, w in enumerate(widths):
        t.columns[i].width = Twips(w)
    return t


def fix_widths(table, widths):
    """Пересчитать tcW всех ячеек с учётом gridSpan после объединений."""
    for row in table.rows:
        pos = 0
        for tc in row._tr.findall(qn('w:tc')):
            tcPr = tc.find(qn('w:tcPr'))
            gs = 1
            gs_el = tcPr.find(qn('w:gridSpan')) if tcPr is not None else None
            if gs_el is not None:
                gs = int(gs_el.get(qn('w:val')))
            w = sum(widths[pos:pos + gs]) if pos < len(widths) else widths[-1]
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.insert(0, tcW)
            tcW.set(qn('w:w'), str(w))
            tcW.set(qn('w:type'), 'dxa')
            pos += gs


def set_tr_heights(table, heights):
    """Высоты строк (twips, atLeast — как в эталоне).
    heights: список по строкам ЛИБО dict {индекс: val} (пропущенные — без высоты)."""
    for i, row in enumerate(table.rows):
        val = heights[i] if isinstance(heights, (list, tuple)) else heights.get(i)
        if val is None:
            continue
        tr = row._tr
        trPr = tr.find(qn('w:trPr'))
        if trPr is None:
            trPr = OxmlElement('w:trPr')
            tr.insert(0, trPr)
        th = trPr.find(qn('w:trHeight'))
        if th is None:
            th = OxmlElement('w:trHeight')
            trPr.append(th)
        th.set(qn('w:val'), str(val))


def _num(v):
    """'-'/''/None -> 0, иначе int."""
    if v in (None, '', '-'):
        return 0
    return int(v)


def _dash0(v):
    """0/'-'/'' -> '-', иначе число строкой."""
    n = _num(v)
    return '-' if n == 0 else str(n)


# ======================================================================
# Титульный лист
# ======================================================================

def ensure_zagolovok1(doc):
    """Стиль Заголовок1 (как в эталоне): Times New Roman 12pt."""
    try:
        st = doc.styles['Заголовок1']
    except KeyError:
        st = doc.styles.add_style('Заголовок1', WD_STYLE_TYPE.PARAGRAPH)
    st.font.name = FONT
    st.font.size = Pt(12)
    return st


def build_title(doc, d):
    ensure_zagolovok1(doc)
    Z1 = 'Заголовок1'
    C = WD_ALIGN_PARAGRAPH.CENTER
    L = WD_ALIGN_PARAGRAPH.LEFT

    # --- Министерство (строго по эталону: БЕЗ «общего и профессионального») ---
    add_p(doc, 'МИНИСТЕРСТВО ОБРАЗОВАНИЯ РОСТОВСКОЙ ОБЛАСТИ', style=Z1, align=C, bold=True)
    add_p(doc, '')
    add_p(doc, 'ГОСУДАРСТВЕННОЕ БЮДЖЕТНОЕ ПРОФЕССИОНАЛЬНОЕ ОБРАЗОВАТЕЛЬНОЕ '
               'УЧРЕЖДЕНИЕ РОСТОВСКОЙ ОБЛАСТИ', style=Z1, align=C, bold=True)
    add_p(doc, '«САЛЬСКИЙ ИНДУСТРИАЛЬНЫЙ ТЕХНИКУМ»', style=Z1, align=C, bold=True)
    add_p(doc, '(ГБПОУ РО «СИТ»)', style=Z1, align=C, bold=True)
    add_p(doc, '')

    # --- Блок УТВЕРЖДАЮ (безрамочная таблица 3x3, текст в 3-й колонке) ---
    UTV_WIDTHS = [4361, 1417, 3793]
    t_utv = new_table(doc, 3, 3, UTV_WIDTHS, bordered=False)
    cell_text(t_utv.cell(0, 2), 'УТВЕРЖДАЮ', bold=True, size=12)
    c = t_utv.cell(1, 2)
    c.text = ''
    p1 = c.paragraphs[0]
    p1.alignment = L
    run = p1.add_run('Зам. директора ГБПОУ РО «СИТ»')
    _fmt_run(run, size=12)
    p2 = c.add_paragraph()
    p2.alignment = L
    run = p2.add_run('______________ ' + d.get('director', 'Т.В. Якимова'))
    _fmt_run(run, size=12)
    c2 = t_utv.cell(2, 2)
    c2.text = ''
    p3 = c2.paragraphs[0]
    p3.alignment = L
    run = p3.add_run('«____» ______________ 20 ____ г.')
    _fmt_run(run, size=12)
    for row in t_utv.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    fix_widths(t_utv, UTV_WIDTHS)

    add_p(doc, '')
    add_p(doc, 'КАЛЕНДАРНО-ТЕМАТИЧЕСКИЙ ПЛАН', style=Z1, align=C, bold=True)
    add_p(doc, '')

    # --- Семестры / группа / ПМ / МДК / профессия (значения подчёркнуты) ---
    add_fillin_p(doc, [
        ('на __', False), (d['semesters'], True), ('__ семестр(ы) _', False),
        (d['year_start'], True), ('___ - __', False), (d['year_end'], True),
        ('___ учебного года _____', False), (d['course'], True), ('______ курс', False),
    ], style=Z1, align=L)
    add_fillin_p(doc, [
        ('учебной группы (учебных групп) ___________ ', False),
        (d['group'], True), ('____________________', False),
    ], style=Z1, align=L)
    add_fillin_p(doc, [
        ('Профессиональный модуль: ', False),
        ('ПМ.%s %s' % (d['pm_code'], d['pm_name']), True),
    ], align=L)
    add_fillin_p(doc, [
        ('Междисциплинарные курсы: ', False),
        ('МДК %s %s' % (d['mdk_code'], d['mdk_name']), True),
    ], align=L)
    add_fillin_p(doc, [
        ('по профессии: ', False), (d['profession'], True),
    ], align=L)
    add_p(doc, '')

    # --- Блок часов ---
    att = d.get('attestation', {})
    if not att and d.get('exam_form'):
        att = {'form': d['exam_form'], 'hours': 2}
    att_form = att.get('form', '')
    att_hours = att.get('hours', '')
    # ЧАСЫ — синхронизация с РП (решение владельца, см. logic.md 4.4):
    # титул и Таблица 2 = строка «МДК 05.02 …» тематического плана РП (раздел 3):
    #   total_hours      = кол. 3 строки МДК (объём МДК, для 05.02 — 112)
    #   practice_hours   = кол. 4 строки МДК («в форме практической подготовки», 60)
    # ЭТА цифра ставится и на титул, и в Таблицу 2.
    # НЕ путать с 96 (= 60 лаб/практ МДК + 36 учебной практики УП.05) — это
    # практическая подготовка раздела целиком по табл. 3.1 РП; практика планируется
    # отдельным документом и на титул КТП по МДК не попадает.
    # Слово «час» в скобках склоняется: 8→часов, 2→часа, 60→часов, 112→часов.
    add_fillin_p(doc, [
        ('Объем образовательной программы: ______', False),
        (str(d['total_hours']), True), ('______________ (%s);' % hour_word(d['total_hours']), False),
    ], style=Z1, align=L)
    add_fillin_p(doc, [
        ('в том числе в форме практической подготовки ___', False),
        (str(d['practice_hours']), True), ('______________(%s);' % hour_word(d['practice_hours']), False),
    ], style=Z1, align=L)
    add_fillin_p(doc, [
        ('Учебная нагрузка во взаимодействии с преподавателем________', False),
        (str(d['total_hours']), True), ('_______________ (%s):' % hour_word(d['total_hours']), False),
    ], style=Z1, align=L)
    add_p(doc, 'из нее:', style=Z1, align=L)
    add_fillin_p(doc, [
        ('теоретическое обучение __', False), (str(d['theory_hours']), True),
        ('___ (%s);                    практические занятия __' % hour_word(d['theory_hours']), False),
        (str(d['practice_hours']), True), ('__ (%s);' % hour_word(d['practice_hours']), False),
    ], style=Z1, align=L)
    add_fillin_p(doc, [
        ('лабораторные занятия ', False), (str(d.get('lab_hours', '') or ''), True),
        ('_______ (%s);                   курсовая работа/проект ' % hour_word(d.get('lab_hours')), False),
        (str(d.get('coursework_hours', '') or ''), True), ('_______ (%s);' % hour_word(d.get('coursework_hours')), False),
    ], style=Z1, align=L)
    add_fillin_p(doc, [
        ('самостоятельная работа ', False), (str(d.get('self_hours', '') or ''), True),
        ('_______ (%s);' % hour_word(d.get('self_hours')), False),
        (' ' * 33, False),  # хвостовое выравнивание как в эталоне
    ], style=Z1, align=L)
    add_fillin_p(doc, [
        ('промежуточная аттестация в форме __ ', False),
        (att_form, True), ('   _', False), (str(att_hours), True), ('__(%s).' % hour_word(att_hours), False),
    ], style=Z1, align=L)
    add_p(doc, ' ' * 84 + '(указать форму)', style=Z1, align=L)
    add_p(doc, '')

    # --- Цикловая комиссия ---
    add_p(doc, 'Составлен в соответствии с рабочей программой ПМ %s, '
               'утверждённой __________________' % d['pm_code'], style=Z1, align=L)
    add_p(doc, ' ' * 14 + '(дата утверждения)', style=Z1, align=L)
    add_p(doc, 'Рассмотрен на заседании цикловой комиссии       %s    дисциплин'
          % d.get('commission', 'технических'), style=Z1, align=L)
    add_p(doc, 'Протокол  №_______от________________20_____ года', style=Z1, align=L)
    add_p(doc, 'Председатель цикловой комиссии ___________________________/%s/'
          % d.get('chair', 'Ткаченко А.Н.'), style=Z1, align=L)
    add_p(doc, '')
    add_p(doc, '')
    add_p(doc, 'г. Сальск', style=Z1, align=C)
    add_p(doc, '%s – %s уч. год' % (d['year_start'], d['year_end']), style=Z1, align=C)


# ======================================================================
# Таблица 1 — Распределение часов (12 колонок)
# ======================================================================

T1_WIDTHS = [1277, 567, 567, 1153, 664, 851, 850, 851, 840, 882, 829, 1216]


def build_table1(doc, d):
    add_p(doc, 'Распределение часов по профессиональному модулю',
          align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    add_p(doc, 'Таблица 1', align=WD_ALIGN_PARAGRAPH.RIGHT)

    sems = d.get('table1_semesters', [])
    # Аттестация по МДК (КДЗ) — отдельной строкой ПОСЛЕ резервных строк
    # (по образцу эталона 05.02: «Компл. дифф. зачет», часы в графе 6,
    #  остальные графы пустые). КДЗ входит в ОБЪЁМ последнего семестра.
    att = d.get('attestation', {})
    if not att and d.get('exam_form'):
        att = {'form': d['exam_form'], 'hours': 2}
    att_hours = _num(att.get('hours', 0))
    has_att_row = att_hours > 0
    reserve = max(0, 12 - 5 - len(sems) - (3 if has_att_row else 2))
    total_rows = 5 + len(sems) + reserve + (3 if has_att_row else 2)  # шапка(4)+номера + семестры + резерв + [КДЗ] + Практика + Всего
    t = new_table(doc, total_rows, 12, T1_WIDTHS)

    # --- Шапка: вертикальные объединения ---
    t.cell(0, 0).merge(t.cell(3, 0))   # МДК
    t.cell(0, 1).merge(t.cell(3, 1))   # Курс
    t.cell(0, 2).merge(t.cell(3, 2))   # Семестр
    # --- Горизонтальные объединения R0 ---
    t.cell(0, 3).merge(t.cell(0, 9))   # Объём времени (7 колонок)
    t.cell(0, 10).merge(t.cell(0, 11))  # Практика (2 колонки)
    # --- R1 ---
    t.cell(1, 3).merge(t.cell(3, 3))   # Объем образовательной программы
    t.cell(1, 4).merge(t.cell(1, 8))   # Учебная нагрузка (5 колонок)
    t.cell(1, 9).merge(t.cell(3, 9))   # Самостоятельная работа
    t.cell(1, 10).merge(t.cell(3, 10))  # Учебная практика
    t.cell(1, 11).merge(t.cell(3, 11))  # Производственная практика
    # --- R2 ---
    t.cell(2, 4).merge(t.cell(3, 4))   # Всего, часов
    t.cell(2, 5).merge(t.cell(2, 8))   # в т.ч. (4 колонки)

    H = WD_ALIGN_PARAGRAPH.CENTER
    cell_text(t.cell(0, 0), 'Междисциплинарный курс (индекс МДК)', bold=True, align=H)
    cell_text(t.cell(0, 1), 'Курс', bold=True, align=H)
    cell_text(t.cell(0, 2), 'Семестр', bold=True, align=H)
    cell_text(t.cell(0, 3), 'Объём времени, отведённый на освоение междисциплинарного курса',
              bold=True, align=H)
    cell_text(t.cell(0, 10), 'Практика ', bold=True, align=H)
    cell_text(t.cell(1, 3), 'Объем образовательной программы', bold=True, align=H)
    cell_text(t.cell(1, 4), 'Учебная нагрузка во взаимодействии с преподавателем',
              bold=True, align=H)
    cell_text(t.cell(1, 9), 'Самостоятельная работа обучающегося, часов', bold=True, align=H)
    cell_text(t.cell(1, 10), 'Учебная\nчасов', bold=True, align=H)
    cell_text(t.cell(1, 11), 'Производственная\nчасов', bold=True, align=H)
    cell_text(t.cell(2, 4), 'Всего, часов', bold=True, align=H)
    cell_text(t.cell(2, 5), 'в т.ч.', bold=True, align=H)
    cell_text(t.cell(3, 5), 'Теоретические занятия', bold=True, align=H)
    cell_text(t.cell(3, 6), 'лабораторные работы, часов', bold=True, align=H)
    cell_text(t.cell(3, 7), 'практические занятия, часов', bold=True, align=H)
    cell_text(t.cell(3, 8), 'Курсовые работы (проекты), часов ', bold=True, align=H)

    # --- Номерная строка (НЕ жирная, по эталону) ---
    for c in range(12):
        cell_text(t.cell(4, c), str(c + 1), bold=False, align=H)

    # --- Строки семестров (все ячейки жирные, по эталону) ---
    mdk = 'МДК %s' % d['mdk_code']
    ri = 5
    for s in sems:
        vals = [mdk, str(s.get('course', d['course'])), str(s['semester']),
                str(s['total']), str(s.get('with_teacher', s['total'])),
                str(s['theory']), _dash0(s.get('lab', '-')),
                str(s['practice']), _dash0(s.get('coursework', '-')),
                _dash0(s.get('self', '-')), _dash0(s.get('uch_practice', '-')),
                _dash0(s.get('proizv_practice', '-'))]
        for c, v in enumerate(vals):
            cell_text(t.cell(ri, c), v, bold=True, align=H)
        ri += 1

    # --- Резервные пустые строки (как в эталоне — 12 строк всего) ---
    for _ in range(reserve):
        for c in range(12):
            cell_text(t.cell(ri, c), '', align=H)
        ri += 1

    # --- Строка «Компл. дифф. зачет» (по образцу эталона: часы в графе 6, остальные пустые) ---
    if has_att_row:
        row = ['Компл. дифф. зачет'] + [''] * 4 + [str(att_hours)] + [''] * 6
        for c, v in enumerate(row):
            cell_text(t.cell(ri, c), v, bold=True, align=H)
        ri += 1

    # --- Строка «Практика» ---
    row = ['Практика'] + ['-'] * 11
    for c, v in enumerate(row):
        cell_text(t.cell(ri, c), v, bold=True, align=H)
    ri += 1

    # --- Строка «Всего» (графа «Теоретические занятия» = Σ семестров + КДЗ;
    #     объём последнего семестра УЖЕ включает КДЗ — правило данных) ---
    sums = {k: sum(_num(s.get(k, '-')) for s in sems)
            for k in ('total', 'with_teacher', 'theory', 'lab', 'practice',
                      'coursework', 'self', 'uch_practice', 'proizv_practice')}
    row = ['Всего', '', '',
           str(sums['total']), str(sums['with_teacher']),
           str(sums['theory'] + att_hours), _dash0(sums['lab']),
           str(sums['practice']), _dash0(sums['coursework']),
           _dash0(sums['self']), _dash0(sums['uch_practice']),
           _dash0(sums['proizv_practice'])]
    for c, v in enumerate(row):
        cell_text(t.cell(ri, c), v, bold=True, align=H)

    # Высоты строк по эталону: строка 0 — без высоты; шапка 219/234/1758;
    # номерная 237; данные (семестры/резерв/КДЗ/Практика/Всего) — 215
    set_tr_heights(t, {1: 219, 2: 234, 3: 1758, 4: 237,
                       **{i: 215 for i in range(5, total_rows)}})
    fix_widths(t, T1_WIDTHS)

    # --- Подпись под Таблицей 1 (строго по эталону) ---
    att_form = att.get('form', '')
    add_p(doc, '(МДК %s %s) – (%s).' % (d['mdk_code'], d['mdk_name'], att_form))


# ======================================================================
# Таблица 2 — Содержание обучения (11 колонок)
# ======================================================================

T2_WIDTHS = [501, 2933, 668, 709, 834, 692, 785, 909, 922, 865, 856]


def build_table2(doc, d):
    add_p(doc, 'Содержание обучения по профессиональному модулю', bold=True)
    add_p(doc, 'Таблица 2')

    topics = d['topics']
    att = d.get('attestation', {})
    if not att and d.get('exam_form'):
        att = {'form': d['exam_form'], 'hours': 2}
    has_att = bool(att.get('form'))

    n_lessons = sum(len(t['lessons']) for t in topics)
    n_rows = 4 + 1 + sum(1 + len(t['lessons']) for t in topics) + (1 if has_att else 0) + 1
    t = new_table(doc, n_rows, 11, T2_WIDTHS)

    # --- Шапка: объединения (порядок важен) ---
    for c in (0, 1, 7, 8, 9, 10):
        t.cell(0, c).merge(t.cell(2, c))   # вертикальные колонки шапки
    for c in (2, 3, 4):
        t.cell(1, c).merge(t.cell(2, c))   # подшапка нагрузки: vMerge R1-R2
    t.cell(0, 2).merge(t.cell(0, 4))       # «Обязательная учебная нагрузка» gs=3
    t.cell(0, 5).merge(t.cell(1, 6))       # «Коды формируемых компетенции» блок 2x2

    C = WD_ALIGN_PARAGRAPH.CENTER
    L = WD_ALIGN_PARAGRAPH.LEFT
    cell_text(t.cell(0, 0), '№ занятия', bold=True, align=C)
    cell_text(t.cell(0, 1), 'Наименование разделов\nпрофессионального модуля,\nтем и занятий по МДК',
              bold=True, align=C)
    cell_text(t.cell(0, 2), 'Обязательная учебная нагрузка', bold=True, align=C)
    cell_text(t.cell(0, 5), 'Коды формируемых компетенции', bold=True, align=C)
    cell_text(t.cell(0, 7), 'Материальное и информационное\nобеспечение занятий\n', bold=True, align=C)
    cell_text(t.cell(0, 8), 'Задания для студентов\n', bold=True, align=C)
    cell_text(t.cell(0, 9), 'Формы и методы контроля', bold=True, align=C)
    cell_text(t.cell(0, 10), 'ФИО преподавателя \n', bold=True, align=C)
    cell_text(t.cell(1, 2), 'Кол-во\nчасов', bold=True, align=C)
    cell_text(t.cell(1, 3), 'В форме практической подготовки', bold=True, align=C)
    cell_text(t.cell(1, 4), 'Вид занятия', bold=True, align=C)
    cell_text(t.cell(2, 5), 'ОК', bold=True, align=C)
    cell_text(t.cell(2, 6), 'ПК', bold=True, align=C)
    # Номерная строка (НЕ жирная, по эталону)
    for c in range(11):
        cell_text(t.cell(3, c), str(c + 1), bold=False, align=C)

    ri = 4
    # --- Строка МДК ---
    cell_text(t.cell(ri, 1), 'МДК %s   %s' % (d['mdk_code'], d['mdk_name']),
              bold=True, align=L)
    cell_text(t.cell(ri, 2), str(d['total_hours']), bold=True, align=C)
    cell_text(t.cell(ri, 3), str(d['practice_hours']), bold=True, align=C)
    cell_text(t.cell(ri, 10), d.get('teacher_name', ''), bold=False, align=C)
    ri += 1

    # --- Темы и занятия ---
    for topic in topics:
        cell_text(t.cell(ri, 1), topic['name'], bold=True, align=L)
        cell_text(t.cell(ri, 2), str(topic['total_hours']), bold=True, align=C)
        if _num(topic.get('practice_hours', 0)):
            cell_text(t.cell(ri, 3), str(topic['practice_hours']), bold=True, align=C)
        ri += 1
        for les in topic['lessons']:
            cell_text(t.cell(ri, 0), str(les['num']), bold=False, align=C)
            cell_text(t.cell(ri, 1), les['name'], bold=False, align=L)
            cell_text(t.cell(ri, 2), str(les['hours']), bold=False, align=C)
            if _num(les.get('practice_prep', 0)):
                cell_text(t.cell(ri, 3), str(les['practice_prep']), bold=False, align=C)
            cell_text(t.cell(ri, 4), les.get('type', ''), bold=False, align=C)
            cell_text(t.cell(ri, 5), les.get('ok', ''), bold=False, align=C)
            cell_text(t.cell(ri, 6), les.get('pk', ''), bold=False, align=C)
            cell_text(t.cell(ri, 7), les.get('equipment', ''), bold=False, align=C)
            cell_text(t.cell(ri, 8), les.get('task', ''), bold=False, align=C)
            cell_text(t.cell(ri, 9), les.get('control', ''), bold=False, align=C)
            ri += 1

    # --- Строка промежуточной аттестации (строго по эталону:
    #     номер, название, 2 часа; ОК/ПК/вид/обеспечение/задания/контроль — ПУСТО) ---
    if has_att:
        row_name = att.get('row_name', att['form'])
        cell_text(t.cell(ri, 0), str(n_lessons + 1), bold=False, align=C)
        cell_text(t.cell(ri, 1), row_name, bold=True, align=L)
        cell_text(t.cell(ri, 2), str(att.get('hours', 2)), bold=True, align=C)
        ri += 1

    # --- Строка «Итого» ---
    cell_text(t.cell(ri, 1), 'Итого:', bold=True, align=L)
    cell_text(t.cell(ri, 2), str(d['total_hours']), bold=True, align=C)
    cell_text(t.cell(ri, 3), str(d['practice_hours']), bold=True, align=C)

    # Высоты строк по эталону: шапка 20/230/1695, все остальные — 20
    set_tr_heights(t, [20, 230, 1695] + [20] * (n_rows - 3))
    fix_widths(t, T2_WIDTHS)


# ======================================================================
# Таблицы 2а, 2б, 2в, 2г
# ======================================================================

T2A_WIDTHS = [1743, 8713]
T2BC_WIDTHS = [913, 4582, 2034, 2998]
T2G_WIDTHS = [748, 2391, 1789, 2551, 3061]


def build_tables_2a_2g(doc, d):
    C = WD_ALIGN_PARAGRAPH.CENTER
    L = WD_ALIGN_PARAGRAPH.LEFT
    J = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ---------- Таблица 2а ----------
    add_p(doc, 'Материально-техническое обеспечение занятий', align=C)
    add_p(doc, 'Таблица 2а', align=WD_ALIGN_PARAGRAPH.RIGHT)

    equipment = d.get('equipment', [])
    t = new_table(doc, len(equipment) + 2, 2, T2A_WIDTHS)
    cell_text(t.cell(0, 0), '№ п/п', bold=False, align=C)
    cell_text(t.cell(0, 1), 'Материально-техническое обеспечение занятий', bold=False, align=C)
    # Номерная строка (жирная, по эталону)
    cell_text(t.cell(1, 0), '1', bold=True, align=C)
    cell_text(t.cell(1, 1), '2', bold=True, align=C)
    for i, eq in enumerate(equipment):
        cell_text(t.cell(i + 2, 0), str(i + 1), bold=False, align=C)
        cell_text(t.cell(i + 2, 1), eq, bold=False, align=L)
    fix_widths(t, T2A_WIDTHS)

    # ---------- Информационное обеспечение (с новой страницы, как в эталоне) ----------
    page_break(doc)
    add_p(doc, 'Информационное обеспечение обучения', align=C, bold=True)

    def lit_table(label, tbl_label, sources, widths, with_url=False):
        add_p(doc, label, bold=True)
        add_p(doc, tbl_label, align=WD_ALIGN_PARAGRAPH.RIGHT)
        n_cols = 5 if with_url else 4
        t = new_table(doc, len(sources) + 1, n_cols, widths)
        headers = ['№ п/п', 'Наименование', 'Автор', 'Издательство, год издания']
        if with_url:
            headers.append('URL (ссылка)')
        for j, h in enumerate(headers):
            cell_text(t.cell(0, j), h, bold=False, align=C)
        for i, src in enumerate(sources):
            cell_text(t.cell(i + 1, 0), src.get('code', ''), bold=False, align=L, size=10)
            cell_text(t.cell(i + 1, 1), src.get('name', ''), bold=False, align=J, size=9)
            cell_text(t.cell(i + 1, 2), src.get('author', ''), bold=False, align=L, size=9)
            cell_text(t.cell(i + 1, 3), src.get('publisher', ''), bold=False, align=L, size=9)
            if with_url:
                cell_text(t.cell(i + 1, 4), src.get('url', ''), bold=False, align=L, size=9)
        fix_widths(t, widths)
        add_p(doc, '')

    lit_table('Основные источники (ОИ):', 'Таблица 2б',
              d.get('sources_basic', []), T2BC_WIDTHS)
    lit_table('Дополнительные источники (ДИ):', 'Таблица 2в',
              d.get('sources_additional', []), T2BC_WIDTHS)
    # ЭИ — отдельная таблица 2г с колонкой URL (в эталоне подпись «2в» — опечатка,
    # по инструкции нумеруется 2г)
    lit_table('Электронные источники (ЭИ):', 'Таблица 2г',
              d.get('sources_electronic', []), T2G_WIDTHS, with_url=True)


# ======================================================================
# Сборка документа
# ======================================================================

def generate_ktp(data, output_path):
    doc = Document()

    # Стиль Normal: Times New Roman 12pt, без интервалов (как в эталоне)
    normal = doc.styles['Normal']
    normal.font.name = FONT
    normal.font.size = Pt(12)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.insert(0, rfonts)
    rfonts.set(qn('w:ascii'), FONT)
    rfonts.set(qn('w:hAnsi'), FONT)
    rfonts.set(qn('w:eastAsia'), FONT)
    rfonts.set(qn('w:cs'), FONT)

    # Страница: A4 портрет, поля по эталону (twips)
    for section in doc.sections:
        section.page_width = Twips(11906)
        section.page_height = Twips(16838)
        section.top_margin = Twips(709)
        section.bottom_margin = Twips(567)
        section.left_margin = Twips(900)
        section.right_margin = Twips(851)

    build_title(doc, data)
    page_break(doc)
    build_table1(doc, data)
    page_break(doc)
    build_table2(doc, data)
    page_break(doc)
    build_tables_2a_2g(doc, data)

    doc.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description='Генерация КТП DOCX строго по эталону СИТ')
    parser.add_argument('--data', required=True, help='JSON-файл с данными КТП')
    parser.add_argument('--output', required=True, help='Выходной файл .docx')
    args = parser.parse_args()
    with open(args.data, 'r', encoding='utf-8') as f:
        data = json.load(f)

    # Контрольные суммы (предупреждения)
    theory = data.get('theory_hours', 0)
    practice = data.get('practice_hours', 0)
    att = data.get('attestation', {})
    if not att and data.get('exam_form'):
        att = {'form': data['exam_form'], 'hours': 2}
    att_h = _num(att.get('hours', 0))
    if theory + practice + att_h != data['total_hours']:
        print('⚠ ВНИМАНИЕ: теория (%s) + практика (%s) + аттестация (%s) != общий объём (%s)'
              % (theory, practice, att_h, data['total_hours']), file=sys.stderr)

    path = generate_ktp(data, args.output)
    print('✅ КТП сохранён: %s' % path)


if __name__ == '__main__':
    main()
