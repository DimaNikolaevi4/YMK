from docx import Document
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.table import CT_Tbl
import copy

SRC = '/home/z/my-project/YMK/KTP/МДК_05_02_КТП.docx'
DST = '/home/z/my-project/YMK/KTP/МДК_05_02_КТП.docx'

# Reference column widths (EMU) from the sample
COL_WIDTHS = [
    318135,   # 0: № занятия
    1862455,  # 1: Наименование
    1403985,  # 2: Кол-во часов
    1403985,  # 3: В форме практической подготовки
    1403985,  # 4: Вид занятия
    937895,   # 5: ОК
    937895,   # 6: ПК
    577215,   # 7: Материальное и информационное обеспечение
    585470,   # 8: Задания для студентов
    600710,   # 9: Формы и методы контроля
    543560,   # 10: ФИО преподавателя
]


def set_cell_text(cell, text, font_name='Times New Roman', font_size=10, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=None):
    """Clear cell and set text with formatting."""
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    
    p = cell.paragraphs[0]
    p.alignment = alignment
    
    # Spacing
    pPr = p._element.get_or_add_pPr()
    spacing = parse_xml('<w:spacing ' + nsdecls('w') + ' w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>')
    pPr.append(spacing)
    
    if size is not None:
        font_size = size
    
    run = p.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    
    # Set East Asian font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml('<w:rFonts ' + nsdecls('w') + '/>')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def set_cell_multiline(cell, lines, font_name='Times New Roman', font_size=10, bold=False, alignment=WD_ALIGN_PARAGRAPH.CENTER, size=None):
    """Clear cell and set multi-line text with line breaks."""
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    
    p = cell.paragraphs[0]
    p.alignment = alignment
    
    pPr = p._element.get_or_add_pPr()
    spacing = parse_xml('<w:spacing ' + nsdecls('w') + ' w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>')
    pPr.append(spacing)
    
    if size is not None:
        font_size = size
    
    for i, line in enumerate(lines):
        if i > 0:
            br = parse_xml('<w:br ' + nsdecls('w') + '/>')
            p._element.append(br)
        run = p.add_run(line)
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml('<w:rFonts ' + nsdecls('w') + '/>')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:eastAsia'), font_name)


def set_cell_vertical_alignment(cell, align='center'):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = parse_xml('<w:tcPr ' + nsdecls('w') + '/>')
        tc.insert(0, tcPr)
    vAlign = tcPr.find(qn('w:vAlign'))
    if vAlign is None:
        vAlign = parse_xml('<w:vAlign ' + nsdecls('w') + '/>')
        tcPr.append(vAlign)
    vAlign.set(qn('w:val'), align)


def set_cell_width(cell, width_emu):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = parse_xml('<w:tcPr ' + nsdecls('w') + '/>')
        tc.insert(0, tcPr)
    tw = tcPr.find(qn('w:tcW'))
    if tw is None:
        tw = parse_xml('<w:tcW ' + nsdecls('w') + '/>')
        tcPr.append(tw)
    tw.set(qn('w:w'), str(width_emu))
    tw.set(qn('w:type'), 'dxa')


def set_cell_margins(cell, top=0, bottom=0, left=40, right=40):
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = parse_xml('<w:tcPr ' + nsdecls('w') + '/>')
        tc.insert(0, tcPr)
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar_xml = '<w:tcMar ' + nsdecls('w') + '>' \
            '<w:top w:w="0" w:type="dxa"/>' \
            '<w:bottom w:w="0" w:type="dxa"/>' \
            '<w:left w:w="40" w:type="dxa"/>' \
            '<w:right w:w="40" w:type="dxa"/>' \
            '</w:tcMar>'
        tcMar = parse_xml(tcMar_xml)
        tcPr.append(tcMar)


def main():
    doc = Document(SRC)
    
    # ====== EXTRACT DATA FROM OLD TABLE 2 ======
    old_table = doc.tables[2]
    old_data = []
    for j, row in enumerate(old_table.rows):
        cells = [c.text.strip().replace('\n', ' | ') for c in row.cells]
        old_data.append(cells)
    
    # Data rows start from index 2 (skip header rows 0,1)
    data_rows = old_data[2:]
    
    # ====== REMOVE OLD TABLE 2 ======
    old_tbl_elem = old_table._tbl
    parent = old_tbl_elem.getparent()
    next_sibling = old_tbl_elem.getnext()
    parent.remove(old_tbl_elem)
    
    # ====== CREATE NEW TABLE USING PYTHON-DOCX API ======
    NUM_COLS = 11
    NUM_DATA_ROWS = len(data_rows)
    TOTAL_ROWS = 4 + NUM_DATA_ROWS  # 3 header rows + 1 number row + data rows
    
    # We need to create the table in the document body
    # Find the element before which to insert (the next_sibling we saved)
    new_table = doc.add_table(rows=TOTAL_ROWS, cols=NUM_COLS)
    new_tbl_elem = new_table._tbl
    
    # Move the table to the correct position (before next_sibling)
    parent.remove(new_tbl_elem)
    if next_sibling is not None:
        parent.insert(list(parent).index(next_sibling), new_tbl_elem)
    else:
        parent.append(new_tbl_elem)
    
    # ====== SET TABLE PROPERTIES ======
    tblPr = new_tbl_elem.find(qn('w:tblPr'))
    
    # Table width
    total_width = sum(COL_WIDTHS)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = parse_xml('<w:tblW ' + nsdecls('w') + '/>')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), str(total_width))
    tblW.set(qn('w:type'), 'dxa')
    
    # Table borders
    borders_xml = '<w:tblBorders ' + nsdecls('w') + '>' \
        '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>' \
        '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>' \
        '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>' \
        '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>' \
        '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>' \
        '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>' \
        '</w:tblBorders>'
    tblBorders = parse_xml(borders_xml)
    tblPr.append(tblBorders)
    
    # Table grid
    tblGrid = new_tbl_elem.find(qn('w:tblGrid'))
    if tblGrid is None:
        tblGrid = parse_xml('<w:tblGrid ' + nsdecls('w') + '/>')
        new_tbl_elem.insert(1, tblGrid)  # after tblPr
    # Remove existing gridCols
    for gc in tblGrid.findall(qn('w:gridCol')):
        tblGrid.remove(gc)
    for w in COL_WIDTHS:
        gridCol = parse_xml('<w:gridCol ' + nsdecls('w') + ' w:w="' + str(w) + '"/>')
        tblGrid.append(gridCol)
    
    # Set column widths and cell properties for all cells
    for r in range(TOTAL_ROWS):
        for c in range(NUM_COLS):
            cell = new_table.cell(r, c)
            set_cell_width(cell, COL_WIDTHS[c])
            set_cell_vertical_alignment(cell, 'center')
            set_cell_margins(cell)
            # Default empty text
            if cell.paragraphs[0].text == '' and len(cell.paragraphs[0].runs) == 0:
                pass  # leave empty
    
    # ====== FILL HEADER ROWS ======
    # Row 0: Top-level headers
    set_cell_text(new_table.cell(0, 0), '№ занятия', bold=True, size=10)
    set_cell_multiline(new_table.cell(0, 1), 
        ['Наименование разделов', 'профессионального модуля,', 'тем и занятий по МДК'], 
        bold=True, size=10)
    set_cell_text(new_table.cell(0, 2), 'Обязательная учебная нагрузка', bold=True, size=10)
    set_cell_text(new_table.cell(0, 5), 'Коды формируемых компетенции', bold=True, size=10)
    set_cell_multiline(new_table.cell(0, 7),
        ['Материальное и информационное', 'обеспечение занятий'],
        bold=True, size=10)
    set_cell_text(new_table.cell(0, 8), 'Задания для студентов', bold=True, size=10)
    set_cell_text(new_table.cell(0, 9), 'Формы и методы контроля', bold=True, size=10)
    set_cell_text(new_table.cell(0, 10), 'ФИО преподавателя', bold=True, size=10)
    
    # Row 1: Second-level headers
    set_cell_multiline(new_table.cell(1, 2), ['Кол-во', 'часов'], bold=True, size=10)
    set_cell_text(new_table.cell(1, 3), 'В форме практической подготовки', bold=True, size=10)
    set_cell_text(new_table.cell(1, 4), 'Вид занятия', bold=True, size=10)
    
    # Row 2: Third-level headers
    set_cell_text(new_table.cell(2, 5), 'ОК', bold=True, size=10)
    set_cell_text(new_table.cell(2, 6), 'ПК', bold=True, size=10)
    
    # Row 3: Number row
    numbers = ['1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11']
    for c, num in enumerate(numbers):
        set_cell_text(new_table.cell(3, c), num, bold=True, size=10)
    
    # ====== APPLY MERGES ======
    # Vertical merges: cols 0,1 merge rows 0-2; cols 7-10 merge rows 0-2
    # cols 2,3,4: col 2 merges rows 1-2, cols 3,4 merge rows 1-2
    # cols 5-6: horizontal merge rows 0-1, vertical merge row 0-1 for the merged cell
    
    # Col 0: vertical merge rows 0-2
    new_table.cell(0, 0).merge(new_table.cell(2, 0))
    
    # Col 1: vertical merge rows 0-2
    new_table.cell(0, 1).merge(new_table.cell(2, 1))
    
    # Col 2: vertical merge rows 1-2 (row 0 is part of horizontal merge for "Обязательная учебная нагрузка")
    # Actually in the reference: row 0 has "Обязательная учебная нагрузка" spanning cols 2-4
    # rows 1-2: col 2 has "Кол-во часов" merged vertically
    # So: merge(0,2)-(0,4) horizontal, then merge(1,2)-(2,2), merge(1,3)-(2,3), merge(1,4)-(2,4)
    
    # First, horizontal merge row 0 cols 2-4 for "Обязательная учебная нагрузка"
    # But we already set the text in (0,2). We need to merge (0,2) with (0,3) and (0,4)
    merged_load = new_table.cell(0, 2).merge(new_table.cell(0, 4))
    
    # Then vertical merges for cols 2,3,4 rows 1-2
    new_table.cell(1, 2).merge(new_table.cell(2, 2))
    new_table.cell(1, 3).merge(new_table.cell(2, 3))
    new_table.cell(1, 4).merge(new_table.cell(2, 4))
    
    # Cols 5-6: rows 0-1 have "Коды формируемых компетенции" (h-merged + v-merged)
    # Row 0: h-merge 5-6, Row 1: h-merge 5-6, then v-merge the merged cells rows 0-1
    merged_comp_0 = new_table.cell(0, 5).merge(new_table.cell(0, 6))
    merged_comp_1 = new_table.cell(1, 5).merge(new_table.cell(1, 6))
    # Now v-merge these
    merged_comp_0.merge(merged_comp_1)
    
    # Col 7: vertical merge rows 0-2
    new_table.cell(0, 7).merge(new_table.cell(2, 7))
    
    # Col 8: vertical merge rows 0-2
    new_table.cell(0, 8).merge(new_table.cell(2, 8))
    
    # Col 9: vertical merge rows 0-2
    new_table.cell(0, 9).merge(new_table.cell(2, 9))
    
    # Col 10: vertical merge rows 0-2
    new_table.cell(0, 10).merge(new_table.cell(2, 10))
    
    # ====== FILL DATA ROWS ======
    for i, data in enumerate(data_rows):
        row_idx = 4 + i
        
        num_pp = data[0]       # old col 0: № п/п
        name = data[1]         # old col 1: Наименование
        hours = data[2]        # old col 2: Кол-во часов
        pract = data[3]        # old col 3: В форме практ. подготовки
        vid = data[4]          # old col 4: Вид учебного занятия
        osnash = data[5]       # old col 5: Оснащение занятий → Мат. и инф. обеспечение
        zadanie = data[6]      # old col 6: Задание для студентов
        prim = data[7]         # old col 7: Примечание → Формы и методы контроля
        
        # Determine if it's a header/section row
        is_header = (num_pp == '' and vid == '')
        is_total = name.startswith('Итого')
        bold = is_header or is_total
        
        # Col 0: № занятия
        set_cell_text(new_table.cell(row_idx, 0), num_pp, bold=bold, size=10)
        
        # Col 1: Наименование (left-aligned for data rows)
        align = WD_ALIGN_PARAGRAPH.LEFT if not is_header and not is_total else WD_ALIGN_PARAGRAPH.LEFT
        set_cell_text(new_table.cell(row_idx, 1), name, bold=bold, size=10, alignment=align)
        
        # Col 2: Кол-во часов
        set_cell_text(new_table.cell(row_idx, 2), hours, bold=bold, size=10)
        
        # Col 3: В форме практической подготовки
        set_cell_text(new_table.cell(row_idx, 3), pract, bold=bold, size=10)
        
        # Col 4: Вид занятия
        set_cell_text(new_table.cell(row_idx, 4), vid, bold=bold, size=10)
        
        # Col 5: ОК (empty)
        set_cell_text(new_table.cell(row_idx, 5), '', bold=bold, size=10)
        
        # Col 6: ПК (empty)
        set_cell_text(new_table.cell(row_idx, 6), '', bold=bold, size=10)
        
        # Col 7: Материальное и информационное обеспечение
        set_cell_text(new_table.cell(row_idx, 7), osnash, bold=bold, size=10)
        
        # Col 8: Задания для студентов
        set_cell_text(new_table.cell(row_idx, 8), zadanie, bold=bold, size=10)
        
        # Col 9: Формы и методы контроля
        set_cell_text(new_table.cell(row_idx, 9), prim, bold=bold, size=10)
        
        # Col 10: ФИО преподавателя (empty)
        set_cell_text(new_table.cell(row_idx, 10), '', bold=bold, size=10)
    
    # ====== SAVE ======
    doc.save(DST)
    print(f'Saved to {DST}')
    print(f'Total data rows: {NUM_DATA_ROWS}')
    print(f'Total table rows: {TOTAL_ROWS}')
    print(f'Actual rows in table: {len(new_table.rows)}')


if __name__ == '__main__':
    main()