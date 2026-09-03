import zipfile
from docx import Document
from docx.oxml.ns import qn

DOC = '/home/z/my-project/YMK/KTP/МДК_05_02_КТП.docx'

z = zipfile.ZipFile(DOC)
bad = z.testzip()
print(f'ZIP: {"OK" if not bad else bad}')

doc = Document(DOC)
errors = 0

# ---- gridSpan check using actual XML cells (not python-docx proxies) ----
def check_gridspan(table, name, expected):
    errs = 0
    for ri, row in enumerate(table.rows):
        tcs = row._tr.findall(qn('w:tc'))
        total = 0
        for tc in tcs:
            tcPr = tc.find(qn('w:tcPr'))
            span = 1
            if tcPr is not None:
                g = tcPr.find(qn('w:gridSpan'))
                if g is not None:
                    span = int(g.get(qn('w:val')))
            total += span
        if total != expected:
            print(f'ERR {name} R[{ri}]: gridSpan sum={total} (expect {expected})')
            errs += 1
    print(f'{name} gridSpan({expected}): {"OK" if errs == 0 else f"{errs} ERRORS"}')
    return errs

errors += check_gridspan(doc.tables[1], 'T1', 12)
errors += check_gridspan(doc.tables[2], 'T2', 11)

# ---- Hours check: only lesson rows (Урок / Практ. занятие / Диф) ----
t2 = doc.tables[2]
th = pr = diff_h = 0
for row in t2.rows:
    h = row.cells[2].text.strip()
    v = row.cells[4].text.strip()
    if not h.isdigit():
        continue
    num = row.cells[0].text.strip()
    if num == '':
        continue  # skip theme headers, МДК row, Итого row
    hours = int(h)
    if v == 'Урок':
        th += hours
    elif v == 'Практ. занятие':
        pr += hours
    elif 'Диф' in row.cells[1].text:
        diff_h += hours

total = th + pr + diff_h
print(f'Lesson hours: total={total} theory={th} practice={pr} diff={diff_h}', end=' ')
if th + pr + diff_h == 112 and th == 52 and pr == 60:
    print('OK')
else:
    print('ERR')
    errors += 1

# Check old text
old_found = sum(1 for r in t2.rows if chr(1047)+chr(1072)+chr(1097)+chr(1080)+chr(1090)+chr(1072)+' пр.работы' in r.cells[9].text)
print(f'Old text remaining: {old_found}', 'OK' if old_found == 0 else 'ERR')
if old_found:
    errors += 1

# Diff count
diff = sum(1 for r in t2.rows if chr(1044)+chr(1080)+chr(1092) in r.cells[1].text)
print(f'Diff rows: {diff}', 'OK' if diff == 1 else 'ERR')
if diff != 1:
    errors += 1

# Table structure
print(f'Total tables: {len(doc.tables)} (expect 7)')
print(f'DI: {len(doc.tables[5].rows)} rows (expect 4)')
print(f'EI: {len(doc.tables[6].rows)} rows, {len(doc.tables[6].columns)} cols (expect 6x5)')

# Bardakov
fio = t2.rows[4].cells[10].text.strip()
print(f'Bardakov: "{fio}"', 'OK' if fio == 'Bardakov' or fio == chr(1041)+chr(1072)+chr(1088)+chr(1076)+chr(1072)+chr(1082)+chr(1086)+chr(1074) else 'ERR')

# Title
print(f'Ministry: "{doc.paragraphs[0].text[:50]}..."')
for p in doc.paragraphs:
    if chr(1050)+chr(1040)+chr(1051)+chr(1045)+chr(1053)+chr(1044)+chr(1040)+chr(1056)+chr(1053) in p.text:
        sz = p.runs[0].font.size
        print(f'KTP size: {sz} (expect 152400=12pt)')
        break

print(f'\nTOTAL ERRORS: {errors}')
